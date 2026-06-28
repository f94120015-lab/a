import os
import unicodedata
from docx import Document

files = [
    "/Users/faruknafizfazlioglu/Desktop/Bölüm dersleri promptu.docx",
    "/Users/faruknafizfazlioglu/Desktop/Academic reading sıralama.docx",
    "/Users/faruknafizfazlioglu/Desktop/dersler düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/İKİNCİ BÖLÜM.docx",
    "/Users/faruknafizfazlioglu/Desktop/eng dosya son halİ.docx",
    "/Users/faruknafizfazlioglu/Desktop/Amok Soru Yapılanlar.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok WORD.docx"
]

target = "When he comes"

for f in files:
    normalized_f = unicodedata.normalize('NFD', f)
    if os.path.exists(normalized_f):
        f = normalized_f
    if not os.path.exists(f):
        print(f"Not found: {f}")
        continue
    try:
        doc = Document(f)
        for i, p in enumerate(doc.paragraphs):
            if target.lower() in p.text.lower():
                print(f"\nFOUND in {os.path.basename(f)} at paragraph {i}:")
                # Print paragraphs around
                for j in range(max(0, i - 2), min(len(doc.paragraphs), i + 200)):
                    txt = doc.paragraphs[j].text.strip()
                    if txt:
                        print(f"P {j}: {txt}")
                break
    except Exception as e:
        print(f"Error reading {f}: {e}")
