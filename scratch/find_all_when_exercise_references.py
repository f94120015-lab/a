import os
import unicodedata
from docx import Document

files = [
    "/Users/faruknafizfazlioglu/Desktop/Bölüm dersleri promptu.docx",
    "/Users/faruknafizfazlioglu/Desktop/Academic reading sıralama.docx",
    "/Users/faruknafizfazlioglu/Desktop/dersler düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/İKİNCİ BÖLÜM.docx",
    "/Users/faruknafizfazlioglu/Desktop/eng dosya son halİ.docx",
    "/Users/faruknafizfazlioglu/Desktop/Amok Soru Yapılanlar.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok WORD.docx"
]

target_headers = ["Alıştırma 136", "Alıştırma 137", "Alıştırma 138", "Alıştırma 139", "Alıştırma 140", "Alıştırma 141", "Alıştırma 142", "Alıştırma 143", "Alıştırma 144"]

for f in files:
    normalized_f = unicodedata.normalize('NFD', f)
    if os.path.exists(normalized_f):
        f = normalized_f
    if not os.path.exists(f):
        continue
    try:
        doc = Document(f)
        found = []
        for i, p in enumerate(doc.paragraphs):
            txt = p.text
            for th in target_headers:
                if th.lower() in txt.lower():
                    found.append((i, txt))
        if found:
            print(f"\n--- FILE: {os.path.basename(f)} ({len(doc.paragraphs)} paragraphs) ---")
            for pos, title in found:
                print(f"  P {pos}: {title}")
    except Exception as e:
        print(f"Error: {e}")
