import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_integration_report():
    doc = Document()
    
    # Page setup / margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AMOK İNGİLİZCE ÖĞRENME UYGULAMASI\nYÜKSEK ÖNCELİKLİ DİLBİLGİSİ ENTEGRASYON RAPORU")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(15)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 117, 89) # Pastel Dark Green

    doc.add_paragraph("\n")

    # Section 1: Overview
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Genel Bakış")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(13)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(46, 117, 89)

    p1 = doc.add_paragraph(
        "Kullanıcının dilbilgisel hedeflerine uyum sağlamak amacıyla uygulamadaki mevcut alıştırma cümleleri "
        "üzerinde gramer odaklı bir denetim gerçekleştirilmiş ve hedeflenen [noun + noun] (İsim + İsim) dilbilgisi kalıbını "
        "tam olarak yansıtan 48 adet özgün örnek cümle uygulamaya yüklenmiştir. Diğer derslerde yapılan geçici düzenlemeler "
        "kullanıcının isteği üzerine geri alınarak orijinal kitap yapısı korunmuştur.\n"
    )
    p1.runs[0].font.name = 'Arial'
    p1.runs[0].font.size = Pt(11)

    # Section 2: Integration Locations
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Entegrasyon Konumları ve Teknik Detaylar")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(13)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(46, 117, 89)

    p2 = doc.add_paragraph(
        "Yapılan tüm kod değişiklikleri data.js veritabanı dosyasında gerçekleştirilmiş olup, "
        "kullanıcı arayüzünde ilgili dersin başlıkları ve alıştırma içerikleri güncellenmiştir. "
        "Aşağıdaki tabloda gramer yapısının nereye entegre edildiği detaylı olarak sunulmuştur:\n"
    )
    p2.runs[0].font.name = 'Arial'
    p2.runs[0].font.size = Pt(11)

    # Integration Location Table
    table_loc = doc.add_table(rows=2, cols=5)
    table_loc.style = 'Light Shading Accent 2'

    headers_loc = ["Dilbilgisi Yapısı", "Arayüz Konumu (Ders)", "Veri Değişkeni (data.js)", "Satır Aralığı", "Soru Sayısı"]
    for i, head in enumerate(headers_loc):
        cell = table_loc.cell(0, i)
        p = cell.paragraphs[0]
        run = p.add_run(head)
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    loc_data = [
        ("[noun + noun]", "Bölüm 10 -> 33. Ders (A. İsim + İsim)", "unit3Lesson10SentencesRaw\nunit3Lesson10Exercises", "18707 - 18780\n19295 - 19334", "48 Soru (İndeks 0-47)")
    ]

    for row_idx, data in enumerate(loc_data):
        row = table_loc.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row[col_idx].text = text

    for row in table_loc.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.5)

    doc.add_paragraph("\n")

    # Section 3: Sample Sentences
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Yüklenen Örneklerden Seçmeler")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(13)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(46, 117, 89)

    p3 = doc.add_paragraph(
        "Aşağıda, uygulamaya yüklenen ve doğrudan isim tamlaması yapısını pekiştiren "
        "bazı örnek cümleler ve Türkçe karşılıkları verilmiştir:\n"
    )
    p3.runs[0].font.name = 'Arial'
    p3.runs[0].font.size = Pt(11)

    # Noun-noun Table
    doc.add_paragraph().add_run("A. [noun + noun] Yapısı Örnekleri (33. Ders)").font.bold = True
    table_nn = doc.add_table(rows=6, cols=3)
    table_nn.style = 'Light Shading Accent 2'
    
    headers_nn = ["İngilizce Cümle", "Türkçe Çeviri", "Hedef Yapı"]
    for i, head in enumerate(headers_nn):
        cell = table_nn.cell(0, i)
        run = cell.paragraphs[0].add_run(head)
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    nn_data = [
        ("We analyzed the database infrastructure.", "Veri tabanı altyapısını analiz ettik.", "database infrastructure"),
        ("They configured the cloud repository.", "Bulut deposunu yapılandırdılar.", "cloud repository"),
        ("The simulation stress was high.", "Simülasyon stresi yüksekti.", "simulation stress"),
        ("We resolved the compile errors.", "Derleme hatalarını çözdük.", "compile errors"),
        ("The network architecture is secure.", "Ağ mimarisi güvenlidir.", "network architecture")
    ]
    for row_idx, data in enumerate(nn_data):
        row = table_nn.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row[col_idx].text = text

    for row in table_nn.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.5)

    doc.add_paragraph("\n")

    # Section 4: Verification and Conclusion
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Doğrulama ve Sonuç")
    h4_run.font.name = 'Arial'
    h4_run.font.size = Pt(13)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(46, 117, 89)

    p4 = doc.add_paragraph(
        "Yapılan tüm güncellemelerin ardından check-lessons.js betiği çalıştırılarak veritabanı "
        "bütünlüğü ve soru sayıları doğrulanmıştır. Yapılan doğrulamada:\n"
    )
    p4.runs[0].font.name = 'Arial'
    p4.runs[0].font.size = Pt(11)

    v1 = doc.add_paragraph(style='List Bullet')
    v1_run = v1.add_run("Tüm modüller, alıştırmalar ve veri havuzları hatasız derlenmiştir.")
    v1_run.font.name = 'Arial'
    v1_run.font.size = Pt(11)

    v2 = doc.add_paragraph(style='List Bullet')
    v2_run = v2.add_run("Uygulama genelinde toplam yüklü soru sayısı 6347 olarak doğrulanmış ve tüm alıştırmaların sağlıklı şekilde yüklendiği onaylanmıştır.")
    v2_run.font.name = 'Arial'
    v2_run.font.size = Pt(11)

    v3 = doc.add_paragraph(style='List Bullet')
    v3_run = v3.add_run("17. Ders ve 20. Ders, kullanıcının talebi doğrultusunda orijinal yapılarına ve soru havuzlarına başarıyla döndürülmüştür.")
    v3_run.font.name = 'Arial'
    v3_run.font.size = Pt(11)

    # Save document
    output_path = "/Users/faruknafizfazlioglu/Desktop/amok_entegrasyon_raporu.docx"
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == "__main__":
    create_integration_report()
