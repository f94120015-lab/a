import os
from docx import Document

f = "/Users/faruknafizfazlioglu/Desktop/amok/Alıstırma_Yazılı_Ceviri_Sınırlandırma_Raporu.docx"
try:
    doc = Document(f)
    print(f"Total paragraphs in {os.path.basename(f)}: {len(doc.paragraphs)}")
    found = False
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        if "swim" in text or "alter" in text or "avoid" in text:
            print(f"Match at {i}: {p.text[:100]}")
            found = True
    if not found:
        print("No matches.")
except Exception as e:
    print(f"Error reading: {e}")
