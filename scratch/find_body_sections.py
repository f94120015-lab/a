import os
from docx import Document
import unicodedata

def find_body_sections():
    f = "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx"
    normalized = unicodedata.normalize('NFC', f)
    doc = Document(normalized)
    
    # We want to search for paragraph occurrences of XVI, XVII, etc. in the document body after the TOC.
    # The TOC ends around paragraph 112. Let's inspect from paragraph 112 onwards.
    for i, p in enumerate(doc.paragraphs[112:], start=112):
        text = p.text.strip()
        if "XVI" in text or "XVII" in text or "XVIII" in text:
            print(f"P {i}: {text[:120]}")
            # print the next 5 paragraphs
            for j in range(1, 6):
                if i + j < len(doc.paragraphs):
                    print(f"  + P {i+j}: {doc.paragraphs[i+j].text.strip()[:120]}")

if __name__ == "__main__":
    find_body_sections()
