import os
from docx import Document

book_path = "/Users/faruknafizfazlioglu/Desktop/Amok kitabı-44444444.docx"
if not os.path.exists(book_path):
    print("Book file not found")
    exit()

doc = Document(book_path)
print(f"Loaded book with {len(doc.paragraphs)} paragraphs.")

search_sentences = [
    "Is the data valid",
    "Are the documents ready",
    "Was the concept clear",
    "Did you analyze it",
    "Why is the data wrong",
    "Why did they analyze it",
    "At which level is it"
]

for target in search_sentences:
    print(f"\nSearching for '{target}'...")
    found = False
    for i, p in enumerate(doc.paragraphs):
        if target.lower() in p.text.lower():
            found = True
            # print surrounding paragraphs to see if translations are there
            start_p = max(0, i - 1)
            end_p = min(len(doc.paragraphs), i + 3)
            for j in range(start_p, end_p):
                print(f"  P {j}: {doc.paragraphs[j].text.strip()}")
    if not found:
        print("  Not found in paragraphs.")
