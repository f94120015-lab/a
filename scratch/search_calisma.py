import docx
import os

files = [
    "/Users/faruknafizfazlioglu/Desktop/Amok Soru Yapılanlar.docx",
    "/Users/faruknafizfazlioglu/Desktop/soru örnekleri bölümler.docx"
]

for filepath in files:
    filename = os.path.basename(filepath)
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        continue
    print(f"\n========================================\nFILE: {filename}")
    try:
        doc = docx.Document(filepath)
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if "17. ders" in text.lower() or "20. ders" in text.lower() or "ders 17" in text.lower() or "ders 20" in text.lower():
                print(f"P {i}: {text}")
                # Print 5 paragraphs before and 10 paragraphs after
                for j in range(max(0, i - 5), min(len(doc.paragraphs), i + 15)):
                    print(f"  [{j}]: {doc.paragraphs[j].text.strip()}")
                print("-" * 50)
    except Exception as e:
        print(f"Error: {e}")
