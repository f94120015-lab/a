from docx import Document

doc = Document("/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx")
for idx in range(6450, min(len(doc.paragraphs), 6900)):
    txt = doc.paragraphs[idx].text.strip()
    if txt.startswith("Alıştırma") or "when" in txt.lower():
        print(f"P {idx}: {txt}")
