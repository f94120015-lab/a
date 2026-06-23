import os
from docx import Document

files = [
    "/Users/faruknafizfazlioglu/Desktop/Bölüm dersleri promptu.docx",
    "/Users/faruknafizfazlioglu/Desktop/Academic reading sıralama.docx",
    "/Users/faruknafizfazlioglu/Desktop/dersler düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/İKİNCİ BÖLÜM.docx",
    "/Users/faruknafizfazlioglu/Desktop/eng dosya son halİ.docx",
    "/Users/faruknafizfazlioglu/Desktop/Amok Soru Yapılanlar.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok WORD.docx"
]

target = "emerging from the study"

for f in files:
    if not os.path.exists(f):
        # Let's try matching with normalized unicode names just in case
        import unicodedata
        normalized_f = unicodedata.normalize('NFD', f)
        if os.path.exists(normalized_f):
            f = normalized_f
        else:
            print(f"File not found: {f}")
            continue
    try:
        doc = Document(f)
        for i, p in enumerate(doc.paragraphs):
            if target.lower() in p.text.lower():
                print(f"\nFOUND in {os.path.basename(f)} at paragraph {i}:")
                # Print from i-2 to i+120 to catch Tip 1, Tip 2, Tip 3
                start = max(0, i - 5)
                end = min(len(doc.paragraphs), i + 150)
                for j in range(start, end):
                    print(f"P {j}: {doc.paragraphs[j].text.strip()}")
                break
    except Exception as e:
        print(f"Error reading {f}: {e}")
