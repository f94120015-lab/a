import os
import glob
from docx import Document

def search():
    files = glob.glob("/Users/faruknafizfazlioglu/Desktop/*.docx")
    for f in files:
        if "~$" in f:
            continue
        try:
            doc = Document(f)
            for i, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                if "bodrum" in text.lower():
                    print(f"FOUND IN: {os.path.basename(f)} at paragraph {i}: {text[:120]}")
                    # Print next 20 paragraphs to see if it's the right list
                    for j in range(1, 40):
                        if i + j < len(doc.paragraphs):
                            print(f"  + P {i+j}: {doc.paragraphs[i+j].text.strip()}")
                    break
        except Exception as e:
            pass

if __name__ == "__main__":
    search()
