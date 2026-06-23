import os, glob
import unicodedata
from docx import Document

desktop_dir = '/Users/faruknafizfazlioglu/Desktop'
docx_files = glob.glob(os.path.join(desktop_dir, '*.docx'))

targets = [
    "emerging from the study",
    "Bölüm 12",
    "30. Ders",
    "31. Ders",
    "32. Ders",
    "Tercüme Kılavuzu",
    "Tip:"
]

for f in docx_files:
    if '~$' in f:
        continue
    f_norm = unicodedata.normalize('NFC', f)
    try:
        doc = Document(f_norm)
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if not text:
                continue
            for t in targets:
                if t.lower() in text.lower():
                    # Print match and surrounding paragraphs
                    print(f"\n--- MATCH '{t}' in {os.path.basename(f_norm)} at P {i} ---")
                    start = max(0, i - 2)
                    end = min(len(doc.paragraphs), i + 25)
                    for j in range(start, end):
                        print(f"P {j}: {doc.paragraphs[j].text.strip()}")
                    print("-----------------------------")
                    break
    except Exception as e:
        # print(f"Error reading {os.path.basename(f_norm)}: {e}")
        pass
