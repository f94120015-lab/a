import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background color of a cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding (margins) of a cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level, space_before=Pt(12), space_after=Pt(6)):
    """Adds a heading with custom spacing and colors."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.keep_with_next = True
    
    # Apply custom colors to heading
    run = p.runs[0]
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(31, 78, 121) # Navy Blue
        # Add a subtle bottom border or spacing
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(90, 90, 90) # Dark Gray
    return p

def create_document():
    # Load comparison data
    with open('scratch/comparison_data.json', 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Title Section
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)
    title_run = title_p.add_run("AMOK DİL ÖĞRENME UYGULAMASI\nBÖLÜM 7 VE BÖLÜM 10 KARŞILAŞTIRMA RAPORU")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)
    
    # Section 1: Genel Bakış
    add_heading_styled(doc, "1. Genel Bakış", level=1)
    
    p = doc.add_paragraph(
        "Bu rapor, AMOK dil öğrenme uygulamasında yer alan Bölüm 7 (Etken/Active Yapısı) ve "
        "Bölüm 10 (Edilgen/Passive Yapısı) ünitelerinin içerdiği soru havuzu yapılarını, cümle adetlerini "
        "ve dilbilgisi modellerini karşılaştırmak amacıyla oluşturulmuştur."
    )
    p.style.font.name = 'Calibri'
    p.style.font.size = Pt(11)
    
    # Table 1: Metrik Karşılaştırma
    table = doc.add_table(rows=7, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["Metrik", "Bölüm 7 (Etken Yapılar)", "Bölüm 10 (Edilgen Yapılar)"]
    row_data = [
        ["Ünite Başlığı", "Özne - Geçişli Fiil + Nesne (Sayfa 32)", "Edilgen (Passive) Yapısı (Sayfa 55)"],
        ["Dilbilgisi Çatısı", "Etken Çatı (Active Voice)", "Edilgen Çatı (Passive Voice)"],
        ["Temel Cümle Yapısı", "Subject + Transitive Verb + Object", "Subject + Be + Past Participle (V3)"],
        ["Ham Cümle Havuzu", "30 Cümle (Bölüm 1 Cümleleri)", "9 Cümle (Bölüm 2 Cümleleri)"],
        ["Toplam Aktif Cümle", "120 Cümle (Ham cümlelerin 4 aşamalı akademik genişletmeleri ile)", "9 Cümle (Yalın edilgen yapılar)"],
        ["Alıştırma / Toplam Soru", "12 Alıştırma Seti / 120 Soru", "1 Alıştırma Seti / 10 Soru"]
    ]
    
    # Header styling
    for col_idx, text in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = text
        set_cell_background(cell, "1F4E79")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        
    # Body rows styling
    for row_idx, row_content in enumerate(row_data):
        for col_idx, text in enumerate(row_content):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if (row_idx + 1) % 2 == 0:
                set_cell_background(cell, "F2F2F2") # Alternating light gray
            p = cell.paragraphs[0]
            if col_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 2: Dilbilgisi ve Yapısal Karşılaştırma
    add_heading_styled(doc, "2. Dilbilgisi ve Yapısal Karşılaştırma", level=1)
    
    p = doc.add_paragraph(
        "Bölüm 7 etken (active) cümle yapılarında eylemi gerçekleştiren özne (Subject) vurgulanırken, "
        "Bölüm 10 edilgen (passive) cümle yapılarında eylemden etkilenen nesne cümlenin öznesi haline gelir ve "
        "yardımcı be fiiliyle birlikte fiilin 3. hali (Past Participle) kullanılır. Aşağıdaki tabloda, iki ünitede "
        "kullanılan ortak fiil köklerinin yapısal karşılaştırmaları yer almaktadır:"
    )
    
    table2 = doc.add_table(rows=8, cols=5)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers2 = ["Fiil Kökü", "Bölüm 7 (Active)", "Türkçe Karşılığı", "Bölüm 10 (Passive)", "Türkçe Karşılığı"]
    row_data2 = [
        ["anticipate", "The sector anticipates growth.", "Sektör büyüme öngörür.", "Growth is anticipated.", "Büyüme öngörülmektedir."],
        ["trigger", "The dynamic triggers reaction.", "Dinamik tepkiyi tetikler.", "The dynamic is triggered.", "Dinamik tetiklenmektedir."],
        ["specify", "The context specifies the criteria.", "Bağlam kriterleri belirler.", "Context is specified.", "Bağlam belirtilmektedir."],
        ["advocate", "Authorities advocate reform.", "Yetkililer reformu savunur.", "Reform is advocated.", "Reform savunulmaktadır."],
        ["define", "The protocol defines parameters.", "Protokol parametreleri tanımlar.", "Parameters are defined.", "Parametreler tanımlanmaktadır."],
        ["calculate", "The script calculates ratios.", "Betik oranları hesaplar.", "Ratios are calculated.", "Oranlar hesaplanmaktadır."],
        ["inspect", "Analysts inspect the framework.", "Analistler çerçeveyi inceler.", "The framework is inspected.", "Çerçeve incelenmektedir."]
    ]
    
    # Header styling table 2
    for col_idx, text in enumerate(headers2):
        cell = table2.cell(0, col_idx)
        cell.text = text
        set_cell_background(cell, "1F4E79")
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.bold = True
        
    # Body rows styling table 2
    for row_idx, row_content in enumerate(row_data2):
        for col_idx, text in enumerate(row_content):
            cell = table2.cell(row_idx + 1, col_idx)
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if (row_idx + 1) % 2 == 0:
                set_cell_background(cell, "F2F2F2")
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 3: Soru Havuzu Analizi
    add_heading_styled(doc, "3. Soru Havuzu Analizi", level=1)
    
    p = doc.add_paragraph(
        "Her iki bölümde de alıştırmalar pedagojik kurallara uygun olarak 10'arlı soru paketleri (exercises) halinde sunulmaktadır. "
        "Alıştırmaların her biri şu 4 farklı soru tipini içermektedir:\n"
        "1. Kelime Eşleştirme (Matching) — İlk 2 soru\n"
        "2. Çoktan Seçmeli Cümle Çevirisi (Multiple-Choice) — 3 ve 4. sorular (Akıllı Çeldirici Algoritması uygulanır)\n"
        "3. Kelime Bankası ile Cümle Kurma (Word-Bank) — 5, 6 ve 7. sorular\n"
        "4. Doğrudan Cümle Çeviri Yazımı (Translation-Text) — 8, 9 ve 10. sorular"
    )
    
    # Show representative Unit 7 Questions
    add_heading_styled(doc, "Bölüm 7 Alıştırma ve Soru Örnekleri (Active)", level=2)
    p_u7 = doc.add_paragraph("Bölüm 7'deki 120 cümlelik genişletilmiş havuzdan üretilen örnek sorular:")
    
    # Add a custom nested list of examples
    ex7_qs = data['unit7']['1']['exercises'][0]['questions']
    for idx, q in enumerate(ex7_qs[:5]):
        p_q = doc.add_paragraph(style='List Bullet')
        p_q.paragraph_format.space_after = Pt(4)
        run_idx = p_q.add_run(f"Soru {idx+1} ({q['type'].upper()}): ")
        run_idx.bold = True
        
        if q['type'] == 'matching':
            pairs_str = ", ".join([f"{p['right']} = {p['left']}" for p in q['pairs']])
            p_q.add_run(f"Kelimeleri Eşleştirin: {pairs_str}")
        elif q['type'] == 'multiple-choice':
            p_q.add_run(f"Soru: {q['prompt']}\nSeçenekler: {', '.join(q['options'])}\nDoğru İndeks: {q['correctIndex']}")
            
    # Show Unit 10 Questions
    add_heading_styled(doc, "Bölüm 10 Alıştırma ve Soru Örnekleri (Passive)", level=2)
    p_u10 = doc.add_paragraph("Bölüm 10'da yer alan ve modulo yöntemiyle 10 soruya tamamlanan edilgen alıştırma soruları:")
    
    ex10_qs = data['unit10']['1']['exercises'][0]['questions']
    for idx, q in enumerate(ex10_qs):
        p_q = doc.add_paragraph(style='List Bullet')
        p_q.paragraph_format.space_after = Pt(4)
        run_idx = p_q.add_run(f"Soru {idx+1} ({q['type'].upper()}): ")
        run_idx.bold = True
        
        if q['type'] == 'matching':
            pairs_str = ", ".join([f"{p['right']} = {p['left']}" for p in q['pairs']])
            p_q.add_run(f"Kelimeleri Eşleştirin: {pairs_str}")
        elif q['type'] == 'multiple-choice':
            p_q.add_run(f"Soru: {q['prompt']}\nSeçenekler: {', '.join(q['options'])}")
        elif q['type'] == 'word-bank':
            p_q.add_run(f"Soru: {q['prompt']} \"{q['enSentence'] if q['isEngToTr'] else q['translation']}\"\nKelimeler: {', '.join(q['words'])}")
        elif q['type'] == 'translation-text':
            p_q.add_run(f"Soru: {q['prompt']}\nDoğru Cevap: \"{q['correctSentence']}\"")
            
    # Save the document
    out_filename = "/Users/faruknafizfazlioglu/Desktop/Bölüm 7 ve Bölüm 10 Karşılaştırma Raporu.docx"
    doc.save(out_filename)
    print(f"Comparison report saved successfully to {out_filename}")

if __name__ == '__main__':
    create_document()
