import os
from docx import Document

fpath = "/Users/faruknafizfazlioglu/Desktop/amok WORD.docx"
if os.path.exists(fpath):
    doc = Document(fpath)
    print(f"Total paragraphs in {fpath}: {len(doc.paragraphs)}")
    
    # We want to print from paragraph 4820 to 5110
    with open("scratch/tips_extracted.txt", "w", encoding="utf-8") as f_out:
        for j in range(4820, min(len(doc.paragraphs), 5110)):
            txt = doc.paragraphs[j].text.strip()
            f_out.write(f"P {j}: {txt}\n")
    print("Extraction completed. Saved to scratch/tips_extracted.txt")
else:
    print("File not found")
