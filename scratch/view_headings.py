import os
from docx import Document
import unicodedata

f = "/Users/faruknafizfazlioglu/Desktop/soru örnekleri bölümler.docx"
normalized = unicodedata.normalize('NFD', f)
if not os.path.exists(normalized):
    normalized = unicodedata.normalize('NFC', f)

try:
    doc = Document(normalized)
    print(f"Total paragraphs in {os.path.basename(normalized)}: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        # If it looks like a heading (short, or starts with BÖLÜM, or has all caps)
        if "BÖLÜM" in text or "Ders" in text or text.isupper() or p.style.name.startswith("Heading"):
            print(f"P {i}: {text[:100]}")
except Exception as e:
    print(f"Error: {e}")
