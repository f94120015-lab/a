import docx
import os

filepath = "/Users/faruknafizfazlioglu/Desktop/Amok Soru Yapılanlar.docx"
doc = docx.Document(filepath)

for i in range(1050, 1150):
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if text:
            print(f"P {i}: {text}")
