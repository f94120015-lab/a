import os
from docx import Document
import unicodedata

f1 = "/Users/faruknafizfazlioglu/Desktop/Kübra Çalışma Dosyası. Diğerleri Sil - Kopya ilk hali.docx"
f2 = "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx"

def inspect_file(file_path, start, end):
    normalized = unicodedata.normalize('NFD', file_path)
    if not os.path.exists(normalized):
        normalized = unicodedata.normalize('NFC', file_path)
    if not os.path.exists(normalized):
        print(f"File not found: {file_path}")
        return
    print(f"\n================ INSPECTING {os.path.basename(normalized)} ==================")
    try:
        doc = Document(normalized)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        for idx in range(start, min(len(doc.paragraphs), end)):
            text = doc.paragraphs[idx].text.strip()
            if text:
                print(f"P {idx}: {text}")
    except Exception as e:
        print(f"Error: {e}")

inspect_file(f1, 1710, 1850)
inspect_file(f2, 5080, 5250)
