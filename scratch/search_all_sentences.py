import os
import glob
from docx import Document

desktop_dir = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_dir, "*.docx"))
docx_files += glob.glob(os.path.join(desktop_dir, "**/*.docx"), recursive=True)

target = "is the data valid"

for fpath in docx_files:
    if "~$" in fpath:
        continue
    try:
        doc = Document(fpath)
        print(f"File: {os.path.basename(fpath)}, paragraphs={len(doc.paragraphs)}")
        for i, p in enumerate(doc.paragraphs):
            if target in p.text.lower():
                print(f"  -> Found at P {i}: {p.text.strip()}")
                # print 3 paragraphs before and after
                for j in range(max(0, i-3), min(len(doc.paragraphs), i+4)):
                    print(f"    P {j}: {doc.paragraphs[j].text.strip()}")
    except Exception as e:
        print(f"  Error reading {os.path.basename(fpath)}: {e}")
