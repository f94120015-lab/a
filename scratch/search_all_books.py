import os
from docx import Document

files = [
    "/Users/faruknafizfazlioglu/Desktop/amok pdf bölünmüş/Amok kitabı-2.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok pdf bölünmüş/Amok kitabı-3.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok WORD.docx"
]

target = "is the data valid"

for fpath in files:
    if not os.path.exists(fpath):
        print(f"File not found: {fpath}")
        continue
    try:
        doc = Document(fpath)
        print(f"\nSearching in {os.path.basename(fpath)}, total paragraphs={len(doc.paragraphs)}")
        for i, p in enumerate(doc.paragraphs):
            if target in p.text.lower():
                print(f"Found at paragraph {i}: {p.text.strip()}")
                for j in range(max(0, i-4), min(len(doc.paragraphs), i+5)):
                    print(f"  P {j}: {doc.paragraphs[j].text.strip()}")
    except Exception as e:
        print(f"Error reading {fpath}: {e}")
