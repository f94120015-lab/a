import os, glob
from docx import Document

desktop_dir = '/Users/faruknafizfazlioglu/Desktop'
docx_files = [f for f in glob.glob(os.path.join(desktop_dir, '*.docx')) if not os.path.basename(f).startswith('~$')]

print(f"Found {len(docx_files)} files:")
for f in docx_files:
    name = os.path.basename(f)
    print(f"\n--- FILE: {name} ---")
    try:
        doc = Document(f)
        print(f"Paragraphs count: {len(doc.paragraphs)}")
        # Print first 10 non-empty paragraphs
        printed = 0
        for i, p in enumerate(doc.paragraphs):
            t = p.text.strip()
            if t:
                print(f"  P {i}: {t[:100]}")
                printed += 1
                if printed >= 10:
                    break
    except Exception as e:
        print(f"  Error: {e}")
