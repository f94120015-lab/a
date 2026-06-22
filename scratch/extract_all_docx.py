from docx import Document

doc = Document("/Users/faruknafizfazlioglu/Desktop/matching.docx")

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}")

# Let's check table cells in the first few tables
for i, table in enumerate(doc.tables[:5]):
    print(f"\nTable {i}:")
    for r_idx, row in enumerate(table.rows[:3]):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row {r_idx}: {cells}")
