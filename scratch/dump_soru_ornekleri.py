import docx

doc = docx.Document("/Users/faruknafizfazlioglu/Desktop/soru örnekleri bölümler.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")
for idx, p in enumerate(doc.paragraphs):
    if "66. Ders" in p.text or "Bölüm 19" in p.text or "Level 1" in p.text or "Level 2" in p.text or "Level 3" in p.text:
        print(f"{idx}: {p.text}")
