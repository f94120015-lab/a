from docx import Document

doc = Document("/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx")
target = "when it dries"
for i, p in enumerate(doc.paragraphs):
    if target.lower() in p.text.lower():
        print(f"P {i}: {p.text.strip()}")
        # print 5 paragraphs before and 10 paragraphs after
        start = max(0, i - 5)
        end = min(len(doc.paragraphs), i + 15)
        for j in range(start, end):
            print(f"  P {j}: {doc.paragraphs[j].text.strip()}")
