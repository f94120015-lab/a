import os, glob
import unicodedata
from docx import Document

desktop_dir = '/Users/faruknafizfazlioglu/Desktop'
docx_files = glob.glob(os.path.join(desktop_dir, '*.docx'))

target = "emerging"

for f in docx_files:
    if '~$' in f:
        continue
    f_norm = unicodedata.normalize('NFC', f)
    try:
        doc = Document(f_norm)
        for i, p in enumerate(doc.paragraphs):
            if target.lower() in p.text.lower():
                print(f"FOUND target in file: {os.path.basename(f_norm)} at paragraph {i}:")
                # Print 20 paragraphs before and 100 after
                start = max(0, i - 15)
                end = min(len(doc.paragraphs), i + 120)
                for j in range(start, end):
                    print(f"P {j}: {doc.paragraphs[j].text.strip()}")
                break
    except Exception as e:
        print(f"Error reading {os.path.basename(f_norm)}: {e}")
