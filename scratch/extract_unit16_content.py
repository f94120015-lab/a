import os
from docx import Document
import unicodedata

fpath = "/Users/faruknafizfazlioglu/Desktop/amok WORD.docx"
f_norm = unicodedata.normalize('NFC', fpath)
if os.path.exists(f_norm):
    doc = Document(f_norm)
    with open("scratch/unit16_extracted.txt", "w", encoding="utf-8") as f_out:
        for j in range(5550, min(len(doc.paragraphs), 5645)):
            txt = doc.paragraphs[j].text.strip()
            f_out.write(f"P {j}: {txt}\n")
    print("Extraction completed. Saved to scratch/unit16_extracted.txt")
else:
    print("File not found")
