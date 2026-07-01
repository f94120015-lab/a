import docx
import glob
import os

desktop_path = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_path, "*.docx"))

target = "highly unstable"

for filepath in docx_files:
    filename = os.path.basename(filepath)
    if "~$" in filename:
        continue
    try:
        doc = docx.Document(filepath)
        for i, p in enumerate(doc.paragraphs):
            if target in p.text.lower():
                print(f"FOUND in {filename} Paragraph {i}: {p.text}")
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text for cell in row.cells)
                if target in row_text.lower():
                    print(f"FOUND in {filename} Table {t_idx} Row {r_idx}: {row_text}")
    except Exception as e:
        print(f"Error {filename}: {e}")
