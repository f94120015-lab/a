import docx

doc = docx.Document("/Users/faruknafizfazlioglu/Desktop/Amok Soru Yapılanlar.docx")
print(f"Total paragraphs: {len(doc.paragraphs)}")
for i in range(3650, min(len(doc.paragraphs), 3950)):
    print(f"{i}: {doc.paragraphs[i].text}")
