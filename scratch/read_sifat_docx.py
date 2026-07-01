import docx

filepath = "/Users/faruknafizfazlioglu/Desktop/sıfat.docx"
doc = docx.Document(filepath)

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"P {i}: {text}")

for t_idx, table in enumerate(doc.tables):
    print(f"\nTABLE {t_idx}:")
    for r_idx, row in enumerate(table.rows):
        row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
        print(f"  Row {r_idx}: {row_text}")
