import os
import unicodedata
from docx import Document

filepath = "/Users/faruknafizfazlioglu/Desktop/soru örnekleri bölümler.docx"
filepath = unicodedata.normalize('NFD', filepath)

if not os.path.exists(filepath):
    print("File not found:", filepath)
    exit(1)

doc = Document(filepath)
print(f"Total paragraphs: {len(doc.paragraphs)}")

start_index = -1
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Bölüm 7:" in text:
        start_index = i
        print(f"Found Bölüm 7 at paragraph {i}: {text}")
        break

if start_index == -1:
    print("Bölüm 7 not found")
    exit(1)

# Print paragraphs from start_index onwards, until we hit "Bölüm 8" or end of document
for i in range(start_index, len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    if i > start_index and ("Bölüm 8" in text or "BÖLÜM 8" in text):
        print(f"\nStopped at paragraph {i} because of: {text}")
        break
    print(f"P {i}: {text}")
