import os
import glob
from docx import Document
import unicodedata

desktop_dir = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_dir, "*.docx"))

targets = ["mikrodan", "wh- clause", "wh-", "clause + infinitive", "wh clause", "bölüm 30", "bolum 30"]

for f in docx_files:
    if "~$" in f:
        continue
    f_norm = unicodedata.normalize('NFD', f)
    if not os.path.exists(f_norm):
        f_norm = unicodedata.normalize('NFC', f)
    if not os.path.exists(f_norm):
        continue
        
    try:
        doc = Document(f_norm)
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            text_lower = text.lower()
            for t in targets:
                if t in text_lower:
                    print(f"Match '{t}' in {os.path.basename(f_norm)} at paragraph {i}: '{text[:120]}'")
                    # Print 10 lines
                    for j in range(max(0, i - 1), min(len(doc.paragraphs), i + 20)):
                        print(f"  P {j}: {doc.paragraphs[j].text.strip()}")
                    print("-" * 50)
                    break
    except Exception as e:
        print(f"Error reading {os.path.basename(f_norm)}: {e}")
