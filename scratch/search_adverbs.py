import docx
import glob
import os

desktop_path = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_path, "*.docx"))

keywords = ["highly", "exceptionally", "completely", "significantly", "deeply", "relatively", "dangerously", "perfectly", "thoroughly"]

for filepath in docx_files:
    filename = os.path.basename(filepath)
    if "~$" in filename:
        continue
    try:
        doc = docx.Document(filepath)
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if any(k in text.lower() for k in keywords):
                if len(text) > 10 and len(text) < 500:
                    print(f"{filename} [P {i}]: {text}")
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                if any(k in row_text.lower() for k in keywords):
                    print(f"{filename} [Table {t_idx}, Row {r_idx}]: {row_text[:200]}")
    except Exception as e:
        pass
