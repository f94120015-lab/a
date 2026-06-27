import os
import glob
from docx import Document
import unicodedata

desktop_dir = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_dir, "*.docx"))

target = "how to alter"

for f in docx_files:
    f_norm = unicodedata.normalize('NFC', f)
    try:
        doc = Document(f)
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if target in text.lower():
                print(f"FOUND in: {os.path.basename(f)} at paragraph {i}")
                # Print 150 paragraphs starting from i
                for j in range(i, min(len(doc.paragraphs), i + 150)):
                    p_text = doc.paragraphs[j].text.strip()
                    if p_text:
                        print(f"P {j}: {p_text}")
                break
    except Exception as e:
         print(f"Error reading {os.path.basename(f)}: {e}")
