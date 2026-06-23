import os
from docx import Document
import unicodedata

fpath = "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx"
f_norm = unicodedata.normalize('NFC', fpath)
if os.path.exists(f_norm):
    doc = Document(f_norm)
    print("Searching in amok düzenleme.docx...")
    found = False
    for i, p in enumerate(doc.paragraphs):
        if "remaining in these areas" in p.text.lower():
            found = True
            print(f"Found in paragraph {i}:")
            for j in range(max(0, i-2), min(len(doc.paragraphs), i+15)):
                print(f"  P {j}: {doc.paragraphs[j].text.strip()}")
            break
    if not found:
        print("Not found in amok düzenleme.docx")
else:
    print("File not found")
