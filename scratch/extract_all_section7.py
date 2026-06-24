import os
import unicodedata
from docx import Document

filepath = "/Users/faruknafizfazlioglu/Desktop/soru örnekleri bölümler.docx"
filepath = unicodedata.normalize('NFD', filepath)

doc = Document(filepath)

start_index = -1
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if "Bölüm 7:" in text:
        start_index = i
        break

if start_index == -1:
    print("Bölüm 7 not found")
    exit(1)

paragraphs_in_section = []
for i in range(start_index, len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    if i > start_index and ("Bölüm 8" in text or "BÖLÜM 8" in text):
        break
    paragraphs_in_section.append(text)

print(f"Total paragraphs in section: {len(paragraphs_in_section)}")
for idx, text in enumerate(paragraphs_in_section):
    if text:
        print(f"{idx}: {text[:100]}")
