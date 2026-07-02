import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = docx.Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Off-black
    
    # Colors
    c_primary = RGBColor(0x1B, 0x36, 0x5D)    # Deep Navy
    c_secondary = RGBColor(0x5C, 0x76, 0x8D)  # Steel Blue
    c_accent = RGBColor(0xD9, 0x53, 0x4F)     # Soft Red/Coral
    c_green = RGBColor(0x4B, 0x8A, 0x4B)      # Forest Green
    
    def set_cell_background(cell, fill_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(m)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Document Header / Cover
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("AMOK UYGULAMASI VE MÜFREDAT KİTABI")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = c_primary
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run("Karşılaştırmalı Müfredat & Eksiklik Analiz Raporu")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(16)
    run_sub.font.italic = True
    run_sub.font.color.rgb = c_secondary
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(48)
    run_date = date_p.add_run("Tarih: 2 Temmuz 2026\nHazırlayan: Antigravity AI Partner")
    run_date.font.size = Pt(11)
    run_date.font.color.rgb = c_secondary
    
    doc.add_paragraph("").paragraph_format.space_after = Pt(24)
    
    # Section 1: Yönetici Özeti
    h1 = doc.add_paragraph()
    r1 = h1.add_run("1. YÖNETİCİ ÖZETİ")
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = c_primary
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Bu rapor, AMOK İngilizce eğitim platformunun veritabanı (data.js) ile fiziksel müfredat kitabının "
        "içindekiler listesi (AMOK İçindekiler.docx) arasındaki kapsam ve yapı uyumluluğunu incelemek amacıyla "
        "hazırlanmıştır. İnceleme sonucunda, kitabın ilk bölümü (Bölüm 1) ile ikinci bölümü (Bölüm 2) "
        "detaylı olarak taranmış ve uygulamadaki eksik konular, dersler, yapısal çakışmalar ve müfredata sonradan "
        "eklenen artı değerler tespit edilmiştir."
    )
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(12)
    
    # Section 2: Eksik ve Hatalı Yapılandırılmış Dersler
    h2 = doc.add_paragraph()
    r2 = h2.add_run("2. EKSİK VE YAPISAL OLARAK ÇAKIŞAN DERSLER")
    r2.font.size = Pt(16)
    r2.font.bold = True
    r2.font.color.rgb = c_primary
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    
    p_intro = doc.add_paragraph(
        "Müfredat kitabında açıkça tanımlanmış olmasına rağmen, uygulamada veri haritalama hataları veya "
        "eksiklikler nedeniyle ulaşılamayan veya hatalı görünen dersler aşağıda listelenmiştir:"
    )
    p_intro.paragraph_format.space_after = Pt(6)
    
    # Bullet points
    bullet1 = doc.add_paragraph(style='List Bullet')
    r = bullet1.add_run("Bölüm 3 (Unit ID 2) - B. Edat Takımı + Edat Takımı (Sayfa 23): ")
    r.font.bold = True
    r = bullet1.add_run(
        "Bu ders müfredatta yer almaktadır ve data.js içerisinde soru havuzları ile alıştırmaları "
        "(unit2Lesson2Exercises) tanımlıdır. Ancak ünitenin aktif ders sayısının (numLessons) 1'e "
        "düşürülmesi ve unitSentencesMap içindeki eşleşmenin kaldırılması nedeniyle arayüzde görünmemektedir. "
        "Şu an Bölüm 3 sadece tek bir ders içermektedir."
    )
    
    bullet2 = doc.add_paragraph(style='List Bullet')
    r = bullet2.add_run("Bölüm 23 - G. Şart - a) if (Sayfa 206): ")
    r.font.bold = True
    r.font.color.rgb = c_accent
    r = bullet2.add_run(
        "Teknik bir ders kimliği (Lesson ID) çakışması nedeniyle kullanıcı bu derse ulaşamamaktadır. "
        "Topic 32 (Bölüm 23) 10 dersten oluşmaktadır ve 8. dersinin ID'si 135'tir. Ancak bir sonraki "
        "Topic 33 (Okuma Parçaları 1) de 135 numaralı ID ile başlamaktadır. Sistemde Okuma Parçaları 1 "
        "daha önce kaydedildiği için, arayüzde 135 ID'li ders 'Genel Okuma Metinleri 1' olarak "
        "görünmekte ve 'G. Şart - a) if' dersi tamamen ezilerek gizlenmektedir."
    )
    
    bullet3 = doc.add_paragraph(style='List Bullet')
    r = bullet3.add_run("Bölüm 15 - B. It + olmak + sıfat + mastar + nesnesi (Sayfa 105): ")
    r.font.bold = True
    r = bullet3.add_run(
        "Müfredatta XIV. ünite altında 'B' başlığıyla tanımlanan bu ders, data.js dosyasında "
        "tanımlanmamış ve ünite ders sayısı (numLessons) 1 olarak bırakılmıştır. Alıştırma ve "
        "soru verisi de veritabanında yer almamaktadır."
    )
    
    # Section 3: Müfredat Kitabında Olup Uygulamada Tamamen Eksik Olan Bölümler
    h3 = doc.add_paragraph()
    r3 = h3.add_run("3. UYGULAMADA TAMAMEN EKSİK OLAN BÜYÜK BÖLÜMLER")
    r3.font.size = Pt(16)
    r3.font.bold = True
    r3.font.color.rgb = c_primary
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Kitap Bölüm Adı'
    hdr_cells[1].text = 'Sayfa No'
    hdr_cells[2].text = 'Uygulamadaki Durum / Analiz'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, '1B365D')
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        
    missing_sections = [
        ("Bilgi Yoklaması I", "68", "Tamamen Eksik. Birinci bölüm konularını kapsayan genel değerlendirme testi uygulamada yer almamaktadır."),
        ("En Mutad Edatlar", "72", "Tamamen Eksik. Edat referans ve çalışma listesi uygulamaya eklenmemiştir."),
        ("Anlama ve Tercüme için Kısa Metinler", "128", "Uyumsuz / Farklı. Kitaptaki orijinal okuma parçaları yerine, uygulamada yapay zeka tarafından üretilmiş genel/teknik bir metin (Okuma Parçaları 1) yer almaktadır."),
        ("Bilgi Yoklaması II", "150", "Tamamen Eksik. Kitabın ilk ana bölümünü kapatan genel seviye kontrol testi uygulamada yer almamaktadır."),
        ("Cetvel I ve Cetvel II", "260-261", "Tamamen Eksik. Kitabın sonundaki özet bağlaç/kural tabloları uygulamada yer almamaktadır."),
        ("Sınıflandırılmış Metinler (Adverbial, Adjectival, Noun Clauses)", "262-290", "Tamamen Eksik. Kitabın ikinci bölüm sonundaki geniş okuma ve analiz metinleri (yaklaşık 30 sayfa) uygulamaya hiç aktarılmamıştır.")
    ]
    
    for name, page, status in missing_sections:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = page
        row_cells[2].text = status
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
            # Set background color for zebra striping
            # let's just make it default white with borders
            
    # Style table borders (will use default Word table grids, docx has built-in styling, but we can set style)
    table.style = 'Light Shading Accent 1'

    doc.add_paragraph("").paragraph_format.space_before = Pt(12)

    # Section 4: Uygulamadaki Ekstra Değer Katılmış Bölümler
    h4 = doc.add_paragraph()
    r4 = h4.add_run("4. KİTAPTA OLMAYIP UYGULAMAYA EKLENEN EK DEĞERLER (EXTRA VALUE)")
    r4.font.size = Pt(16)
    r4.font.bold = True
    r4.font.color.rgb = c_primary
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)
    
    p_extra = doc.add_paragraph(
        "Müfredat kitabında yer almamasına rağmen, uygulamanın zenginleştirilmesi amacıyla "
        "sonradan eklenen ve öğrencilere ek pratik imkanı sağlayan konular şunlardır:"
    )
    p_extra.paragraph_format.space_after = Pt(6)
    
    extras = [
        ("Ara Bölüm 2: Tercih Bildiren Yapılar (Would rather)", "4 Ders. Farklı zamanlarda 'would rather' kullanım yapıları ve özne değişim durumları."),
        ("Ara Bölüm 3: Rica ve İzin İsteme Yapıları", "4 Ders. Günlük dilde rica, kibar soru kalıpları ve izin isteme modelleri."),
        ("Ara Bölüm 4: Before Bağlaçlı Cümle Yapıları", "4 Ders. Before bağlacı ile aktif, pasif ve past perfect cümle dizilimlerinin detaylı pratikleri."),
        ("Zaman Zarfları ve Zaman Uyumu (Ekstra Pratik)", "3 Ders. Akademik dilde zaman zarfları ve tense geçiş kuralları."),
        ("Zaman Uyumu: By the time, Since, It is (high) time", "4 Ders. By the time, since, it is high time kalıpları ve sıfat derecelendirmesi + present perfect."),
        ("Öbeksel Kipler (Phrasal Modals) (Ekstra Pratik)", "16 Ders. be used to, be accustomed to, be likely to gibi akademik dilde en çok karşılaşılan öbeksel modal yapıları ve karma test.")
    ]
    
    for ext_title, ext_desc in extras:
        bullet = doc.add_paragraph(style='List Bullet')
        r = bullet.add_run(ext_title + ": ")
        r.font.bold = True
        r.font.color.rgb = c_green
        bullet.add_run(ext_desc)

    # Section 5: Teknik Öneriler ve Çözüm Planı
    h5 = doc.add_paragraph()
    r5 = h5.add_run("5. TEKNİK ÖNERİLER VE ÇÖZÜM YOL HARİTASI")
    r5.font.size = Pt(16)
    r5.font.bold = True
    r5.font.color.rgb = c_primary
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)
    
    rec_intro = doc.add_paragraph(
        "Uygulamanın müfredat kitabı ile %100 uyumlu hale gelmesi ve mevcut hataların giderilmesi için "
        "yapılması gereken teknik çalışmalar şunlardır:"
    )
    rec_intro.paragraph_format.space_after = Pt(6)
    
    rec1 = doc.add_paragraph(style='List Number')
    r = rec1.add_run("ID Çakışmasının Çözülmesi (Ders 135): ")
    r.font.bold = True
    r = rec1.add_run(
        "Topic 33 (Okuma Parçaları 1) ve sonrasındaki tüm derslerin startLessonId değerleri "
        "kaydırılmalıdır. Örneğin Okuma Parçaları 1 startLessonId'si 135 yerine 200 yapılabilir. "
        "Böylece Bölüm 23'ün 8. dersi olan 'G. Şart - a) if' dersi (ID 135) ezilmekten kurtulacak ve dashboard'da aktifleşecektir."
    )
    
    rec2 = doc.add_paragraph(style='List Number')
    r = rec2.add_run("Bölüm 3'ün Eksik Dersinin Geri Kazanılması: ")
    r.font.bold = True
    r = rec2.add_run(
        "data.js'te Topic 2 (Bölüm 3) numLessons değeri 2 yapılmalı ve unitSentencesMap[2] altına "
        "2: unit2Lesson2Exercises eşlemesi yeniden eklenmelidir. Soru verisi zaten mevcut olduğundan "
        "bu değişiklik anında dersi görünür kılacaktır."
    )
    
    rec3 = doc.add_paragraph(style='List Number')
    r = rec3.add_run("Eksik Cümle Verilerinin Girişi: ")
    r.font.bold = True
    r = rec3.add_run(
        "Bölüm 15'in B dersi için soru ve çeviri cümleleri hazırlanıp veritabanına entegre edilmelidir."
    )
    
    rec4 = doc.add_paragraph(style='List Number')
    r = rec4.add_run("Büyük Okuma Bölümleri ve Tabloların Eklenmesi: ")
    r.font.bold = True
    r = rec4.add_run(
        "Bilgi yoklaması testleri, Cetveller ve Sınıflandırılmış Metinler için yeni ünite tanımları "
        "yapılmalı ve içerikleri (yaklaşık 1000 yeni soru/çeviri) veritabanına eklenmelidir."
    )
    
    # Save document
    output_path = "/Users/faruknafizfazlioglu/Desktop/AMOK Müfredat Analizi.docx"
    doc.save(output_path)
    print(f"Report successfully saved to: {output_path}")

if __name__ == "__main__":
    create_report()
