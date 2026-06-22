from docx import Document

doc = Document("/Users/faruknafizfazlioglu/Desktop/Amok kitabı-44444444.docx")

print(f"Total paragraphs in book: {len(doc.paragraphs)}")

start = max(0, 14502 - 100)
end = min(len(doc.paragraphs), 14502 + 100)

for i in range(start, end):
    txt = doc.paragraphs[i].text.strip()
    if txt:
        print(f"P {i}: {txt}")
