import os
from docx import Document

files = [
    "/Users/faruknafizfazlioglu/Desktop/matching.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/amok WORD.docx",
    "/Users/faruknafizfazlioglu/Desktop/dersler düzenleme.docx",
    "/Users/faruknafizfazlioglu/Desktop/Amok kitabı-44444444.docx"
]

target = "Is the data valid"

for fpath in files:
    if not os.path.exists(fpath):
        print(f"File not found: {fpath}")
        continue
    try:
        doc = Document(fpath)
        print(f"\nSearching in: {fpath}")
        
        # Search paragraphs
        for i, p in enumerate(doc.paragraphs):
            if target.lower() in p.text.lower():
                print(f"Paragraph {i}: {p.text}")
                
        # Search tables
        for i, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = [cell.text for cell in row.cells]
                row_str = " | ".join(row_text)
                if target.lower() in row_str.lower():
                    print(f"Table {i}, Row {row_idx}: {row_str}")
    except Exception as e:
        print(f"Error reading {fpath}: {e}")
