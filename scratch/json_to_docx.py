import json
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    # Set shading XML element
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    tblPr = table._tbl.tblPr
    borders_xml = (
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="EAEAEA"/>'
        '<w:insideV w:val="none"/>'
        '</w:tblBorders>'
    )
    tblPr.append(parse_xml(borders_xml))

# Load JSON data
with open("scratch/first_4_units.json", "r", encoding="utf-8") as f:
    data = json.load(f)

doc = Document()

# Page Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Document Title Block
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("AMOK İNGİLİZCE EĞİTİM UYGULAMASI")
run_title.font.name = 'Calibri'
run_title.font.size = Pt(22)
run_title.font.bold = True
run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Premium Navy

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("İlk 4 Bölüm Soru ve Çeviri Havuzu")
run_sub.font.name = 'Calibri'
run_sub.font.size = Pt(14)
run_sub.font.italic = True
run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59) # Slate Grey

doc.add_paragraph().paragraph_format.space_after = Pt(20)

for unit_title, lessons_data in data.items():
    # Unit Header (Heading 1)
    h_unit = doc.add_heading(level=1)
    run_uh = h_unit.add_run(unit_title)
    run_uh.font.name = 'Calibri'
    run_uh.font.size = Pt(16)
    run_uh.font.bold = True
    run_uh.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    h_unit.paragraph_format.space_before = Pt(24)
    h_unit.paragraph_format.space_after = Pt(12)
    h_unit.paragraph_format.keep_with_next = True
    
    for lesson_title, sentences in lessons_data.items():
        if not sentences:
            continue
            
        # Lesson Header (Heading 2)
        h_lesson = doc.add_heading(level=2)
        run_lh = h_lesson.add_run(lesson_title)
        run_lh.font.name = 'Calibri'
        run_lh.font.size = Pt(12)
        run_lh.font.bold = True
        run_lh.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5) # Steel Blue
        h_lesson.paragraph_format.space_before = Pt(14)
        h_lesson.paragraph_format.space_after = Pt(6)
        h_lesson.paragraph_format.keep_with_next = True
        
        # Create Table
        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        add_table_borders(table)
        
        # Set Widths: Column 0 (Num) = 0.5 inches, Column 1 (EN) = 3.0 inches, Column 2 (TR) = 3.0 inches
        col_widths = [Inches(0.5), Inches(3.0), Inches(3.0)]
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "No"
        hdr_cells[1].text = "İngilizce Cümle"
        hdr_cells[2].text = "Türkçe Çeviri"
        
        # Style Header Row
        for idx, cell in enumerate(hdr_cells):
            cell.width = col_widths[idx]
            set_cell_background(cell, "1F4E79") # Premium Navy
            set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
            
            # Format text inside header
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text
            
        # Data Rows
        for s_idx, s in enumerate(sentences):
            row_cells = table.add_row().cells
            
            # Sentence number, English, Turkish
            row_cells[0].text = str(s_idx + 1)
            row_cells[1].text = s.get('en', '')
            row_cells[2].text = s.get('tr', '')
            
            # Zebra striping background color
            bg_color = "F2F2F2" if s_idx % 2 == 1 else "FFFFFF"
            
            for idx, cell in enumerate(row_cells):
                cell.width = col_widths[idx]
                if bg_color != "FFFFFF":
                    set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                
                # Format text
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                if p.runs:
                    run = p.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    # For English sentence, let's make it dark blue, Turkish standard dark
                    if idx == 1:
                        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                    else:
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

# Save Document
out_path = "/Users/faruknafizfazlioglu/Desktop/Amok İlk 4 Bölüm Soru Havuzu.docx"
doc.save(out_path)
print(f"Word document saved to: {out_path}")
