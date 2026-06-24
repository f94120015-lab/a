import os
import glob
import unicodedata
from docx import Document

desktop_path = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_path, "*.docx"))

target = "The data contradicts the theory"

for f in docx_files:
    f_norm = unicodedata.normalize('NFC', f)
    # also normalized to NFD just in case
    f_nfd = unicodedata.normalize('NFD', f)
    path_to_use = f
    if os.path.exists(f_norm):
        path_to_use = f_norm
    elif os.path.exists(f_nfd):
        path_to_use = f_nfd
        
    try:
        doc = Document(path_to_use)
        for i, p in enumerate(doc.paragraphs):
            p_text = p.text.strip()
            if target.lower() in p_text.lower():
                print(f"FOUND target in {os.path.basename(f)} at paragraph {i}:")
                # print 10 paragraphs around it
                start = max(0, i - 2)
                end = min(len(doc.paragraphs), i + 20)
                for j in range(start, end):
                    print(f"  P {j}: {doc.paragraphs[j].text.strip()}")
                break
    except Exception as e:
        pass
