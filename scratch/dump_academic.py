import os
from docx import Document
import unicodedata

f = "/Users/faruknafizfazlioglu/Desktop/Academic reading sıralama.docx"
normalized = unicodedata.normalize('NFD', f)
if not os.path.exists(normalized):
    normalized = unicodedata.normalize('NFC', f)

try:
    doc = Document(normalized)
    print(f"Total paragraphs in {os.path.basename(normalized)}: {len(doc.paragraphs)}")
    with open("scratch/academic_paragraphs.txt", "w", encoding="utf-8") as out:
        for idx in range(100):
            if idx < len(doc.paragraphs):
                text = doc.paragraphs[idx].text.strip()
                out.write(f"P {idx}: {text}\n")
    print("Done writing to scratch/academic_paragraphs.txt")
except Exception as e:
    print(f"Error: {e}")
