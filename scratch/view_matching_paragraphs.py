from docx import Document

doc = Document("/Users/faruknafizfazlioglu/Desktop/matching.docx")
for idx in range(30, 100):
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        txt = p.text.strip()
        print(f"P {idx}: {txt}")
