import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_toc_document():
    doc = Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("AMOK\nİÇİNDEKİLER")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(31, 78, 121) # Navy Blue
    
    doc.add_paragraph("\n")

    # Data structure for T.O.C.
    toc_data = [
        # Part 1
        ("I. BÖLÜM", "header_part"),
        ("I. A. Özne + olmak + isim", "1", "level_1"),
        ("B. Özne + olmak + sıfat", "2", "level_2"),
        ("C. Özne + olmak + sıfat + isim", "4", "level_2"),
        ("D. Özne + olmak + edat takımı", "6", "level_2"),
        
        ("II. A. İsim + edat takımı", "7", "level_1"),
        ("B. İsim + of the + isim", "13", "level_2"),
        ("C. Zamir + of the + isim", "14", "level_2"),
        ("D. İsim + of + isim", "16", "level_2"),
        ("E. İsim + from + isim", "17", "level_2"),
        ("F. İsim + edat takımı + edat takımı", "19", "level_2"),
        
        ("III. A. Fiil + edat takımı", "21", "level_1"),
        ("B. Edat takımı + edat takımı", "23", "level_2"),
        
        ("IV. Özne + geçimli fiil + nesne", "32", "level_1"),
        
        ("V. A. There + olmak + isim", "38", "level_1"),
        ("B. There + olmak + sıfat + isim", "39", "level_2"),
        
        ("VI. Soru strüktürler", "49", "level_1_header"),
        ("A. a) Soru kelimesi + fiil + özne", "49", "level_2"),
        ("b) Soru kelimesi + do + özne + mastar", "50", "level_2"),
        ("B. a) Fiil + özne", "51", "level_2"),
        ("b) Do, does, did + özne + mastar", "52", "level_2"),
        ("C. Edat + soru kelimesi + fiil + özne + mastar", "53", "level_2"),
        
        ("VII. Edilgen strüktürü", "55", "level_1"),
        ("VIII. Edilgen mastarı", "63", "level_1"),
        ("Bilgi yoklaması I", "68", "level_special"),
        ("En mutad edatlar", "72", "level_special"),
        
        ("IX. İsim tamlaması", "72", "level_1_header"),
        ("A. İsim + isim", "72", "level_2"),
        ("B. İsim / Sıfat + isim + isim", "78", "level_2"),
        
        ("X. Present participle sıfatı", "81", "level_1_header"),
        ("A. .........ing + isim", "81", "level_2"),
        ("B. İsim + .........ing + isim", "84", "level_2"),
        
        ("XI. Past participle sıfatı", "85", "level_1_header"),
        ("A. .........ed + isim", "85", "level_2"),
        ("B. Zarf + past participle + isim", "88", "level_2"),
        
        ("XII. Participle takımları", "88", "level_1_header"),
        ("A. İsim + present participle + edat takımı", "88", "level_2"),
        ("B. İsim + present participle + nesne (isim)", "90", "level_2"),
        ("C. İsim + past participle + edat takımı", "95", "level_2"),
        
        ("XIII. Mastar", "98", "level_1_header"),
        ("A. Fiil + mastar", "98", "level_2"),
        ("B. Fiil + mastar + nesnesi", "100", "level_2"),
        
        ("XIV. Strüktürel 'It' cümlesinin öznesi", "103", "level_1_header"),
        ("A. It + olmak + sıfat + mastar", "103", "level_2"),
        ("B. It + olmak + sıfat + mastar + nesnesi", "105", "level_2"),
        
        ("XV. A. Maksad için kullanılan mastar", "107", "level_1"),
        ("B. Maksad için kullanılan mastarın nesnesi", "110", "level_2"),
        
        ("XVI. Fiil ismi + nesnesi (özne olarak)", "112", "level_1_header"),
        ("A. .........ing + isim", "112", "level_2"),
        ("B. Edattan sonra gelen fiil (+ nesnesi)", "113", "level_2"),
        
        ("XVII. A. on/by/in/without + .........ing + nesne", "116", "level_1"),
        ("B. when/while/before/after/since ......ing (+ nesnesi)", "118", "level_2"),
        ("C. when/if/unless/although/until/as/where + past participle", "120", "level_2"),
        
        ("XVIII. Soru kelimesinden sonra gelen mastar", "124", "level_1_header"),
        ("Faydalı mutad cümle takımları", "127", "level_special"),
        ("Anlama ve tercüme için kısa metinler", "128", "level_special"),
        ("Bilgi yoklaması II", "150", "level_special"),
        
        # Part 2
        ("II. BÖLÜM", "header_part"),
        ("I. Zarf cümleciği", "164", "level_1_header"),
        ("A. Zaman", "164", "level_2_header"),
        ("a) before", "164", "level_3"),
        ("b) after", "169", "level_3"),
        ("c) when", "172", "level_3"),
        ("d) since", "181", "level_3"),
        ("e) while/as", "182", "level_3"),
        ("f) until", "185", "level_3"),
        
        ("B. Sebeb", "187", "level_2_header"),
        ("a) because", "187", "level_3"),
        ("b) since", "191", "level_3"),
        ("c) as", "192", "level_3"),
        
        ("C. 'Although'", "194", "level_2"),
        ("D. Derece", "198", "level_2"),
        ("E. Maksat", "201", "level_2"),
        ("F. Netice", "204", "level_2"),
        
        ("G. Şart", "206", "level_2_header"),
        ("a) if", "206", "level_3"),
        ("b) unless", "213", "level_3"),
        
        ("II. Mukayese strüktürleri", "216", "level_1_header"),
        ("A. Basit", "216", "level_2"),
        ("B. 'than'", "218", "level_2"),
        ("C. 'as......as, the same......as'", "222", "level_2"),
        
        ("III. Sıfat cümleciği", "224", "level_1_header"),
        ("a) Who", "224", "level_2"),
        ("b) Which", "227", "level_2"),
        ("c) Where", "231", "level_2"),
        ("d) Whom", "233", "level_2"),
        ("e) Edat ile başlayan", "234", "level_2"),
        ("f) Whose", "236", "level_2"),
        ("g) İşaret kelimesi olmayan", "239", "level_2"),
        
        ("IV. İsim cümleciği", "240", "level_1_header"),
        ("a) Soru kelimesi", "241", "level_2"),
        ("b) <that>", "244", "level_2"),
        ("c) <if/whether>", "247", "level_2"),
        ("d) <the fact that>", "248", "level_2"),
        
        ("V. It + to be + sıfat/past part. + that", "251", "level_1"),
        ("VI. It + to be + sıfat/past part. + mastar + that", "254", "level_1"),
        
        ("VII. Neden ve etkisi strüktürleri", "256", "level_1_header"),
        ("Cetvel I", "260", "level_special"),
        ("Cetvel II", "261", "level_special"),
        ("Sınıflandırılmış metinler", "262", "level_special"),
        ("Adverbial clauses only", "262", "level_special_sub"),
        ("Adjectival clauses only", "268", "level_special_sub"),
        ("Adjectival and Adverbial clauses", "273", "level_special_sub"),
        ("Noun clauses only", "287", "level_special_sub")
    ]

    for item in toc_data:
        if len(item) == 2: # Part headers
            title_text, style_type = item
            page_text = ""
        else:
            title_text, page_text, style_type = item
            
        p = doc.add_paragraph()
        
        # Apply indentation based on style_type
        p_format = p.paragraph_format
        if style_type == "header_part":
            p_format.left_indent = Inches(0)
            p_format.space_before = Pt(14)
            p_format.space_after = Pt(6)
        elif style_type == "level_1_header" or style_type == "level_1":
            p_format.left_indent = Inches(0.2)
            p_format.space_before = Pt(6)
            p_format.space_after = Pt(3)
        elif style_type == "level_2_header" or style_type == "level_2" or style_type == "level_special":
            p_format.left_indent = Inches(0.4)
            p_format.space_before = Pt(3)
            p_format.space_after = Pt(2)
        elif style_type == "level_3" or style_type == "level_special_sub":
            p_format.left_indent = Inches(0.6)
            p_format.space_before = Pt(2)
            p_format.space_after = Pt(1)

        # Style font
        run_title = p.add_run(title_text)
        run_title.font.name = 'Arial'
        
        if style_type == "header_part":
            run_title.font.size = Pt(14)
            run_title.font.bold = True
            run_title.font.color.rgb = RGBColor(31, 78, 121) # Navy Blue
        elif style_type.endswith("header"):
            run_title.font.size = Pt(11.5)
            run_title.font.bold = True
            run_title.font.color.rgb = RGBColor(60, 60, 60)
        else:
            run_title.font.size = Pt(10.5)
            run_title.font.bold = False
            
        if page_text:
            # We want page numbers aligned to the right or separated cleanly.
            # Let's add tab formatting, or dotted leaders.
            # Dotted leaders look great in Word. We can simulate it or just use tabs/spaces.
            # Let's add a tab stop at 6.0 inches right-aligned, or we can just append dot spaces.
            # Dotted leaders:
            dots_count = 80 - len(title_text) - len(page_text)
            if dots_count < 3:
                dots_count = 3
            dots = " ." * (dots_count // 2) + " "
            
            run_dots = p.add_run(dots)
            run_dots.font.name = 'Arial'
            run_dots.font.size = Pt(9.5)
            run_dots.font.color.rgb = RGBColor(180, 180, 180) # Light grey for dots
            
            run_page = p.add_run(page_text)
            run_page.font.name = 'Arial'
            run_page.font.size = Pt(10.5)
            run_page.font.bold = True
            run_page.font.color.rgb = RGBColor(80, 80, 80)

    # Save Word Document
    dest_path = "/Users/faruknafizfazlioglu/Desktop/AMOK İçindekiler.docx"
    doc.save(dest_path)
    print(f"Created file: {dest_path}")

if __name__ == "__main__":
    create_toc_document()
