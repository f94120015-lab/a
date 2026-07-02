import docx

doc = docx.Document("AMOK İçindekiler.docx")
print("Number of paragraphs:", len(doc.paragraphs))
print("Number of tables:", len(doc.tables))

with open("scratch/toc_text.txt", "w", encoding="utf-8") as f:
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            f.write(f"P{i}: {p.text}\n")
    
    for i, table in enumerate(doc.tables):
        f.write(f"\n--- Table {i} ---\n")
        for row in table.rows:
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            f.write(" | ".join(row_text) + "\n")

print("Done writing scratch/toc_text.txt")
