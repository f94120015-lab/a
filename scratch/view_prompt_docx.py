import os
from docx import Document

f = "/Users/faruknafizfazlioglu/Desktop/Bölüm dersleri promptu.docx"
try:
    doc = Document(f)
    print(f"Total paragraphs in {os.path.basename(f)}: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"P {i}: {text}")
except Exception as e:
    print(f"Error reading: {e}")
