import os
from docx import Document
import unicodedata

fpath = "/Users/faruknafizfazlioglu/Desktop/Selma İngiliş Eng.docx"
f_norm = unicodedata.normalize('NFC', fpath)
if os.path.exists(f_norm):
    doc = Document(f_norm)
    print(f"Total paragraphs in {f_norm}: {len(doc.paragraphs)}")
    for i, p in enumerate(doc.paragraphs[:150]):
        txt = p.text.strip()
        if txt:
            print(f"P {i}: {txt[:100]}")
else:
    print("File not found")
