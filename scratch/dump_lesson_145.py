from docx import Document

doc = Document("/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx")
for idx in range(6460, min(len(doc.paragraphs), 6605)):
    txt = doc.paragraphs[idx].text.strip()
    if txt:
        print(f"P {idx}: {txt}")
