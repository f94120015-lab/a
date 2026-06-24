import os
import glob
import unicodedata
from docx import Document

desktop_path = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_path, "*.docx"))

print(f"Found docx files: {[os.path.basename(f) for f in docx_files]}")

for f in docx_files:
    f_norm = unicodedata.normalize('NFC', f)
    # also normalized to NFD just in case
    f_nfd = unicodedata.normalize('NFD', f)
    path_to_use = f
    if os.path.exists(f_norm):
        path_to_use = f_norm
    elif os.path.exists(f_nfd):
        path_to_use = f_nfd
        
    try:
        doc = Document(path_to_use)
        text_content = []
        for p in doc.paragraphs:
            text_content.append(p.text)
        full_text = "\n".join(text_content)
        
        # Check if the document mentions Bölüm 7, Ders 20, or the translation guide text
        targets = ["Bölüm 7", "Ders 20", "20. Ders", "Akademik makalelerde"]
        found_targets = [t for t in targets if t.lower() in full_text.lower()]
        
        if found_targets:
            print(f"\n--- {os.path.basename(f)} matches {found_targets} ---")
            # Let's find exactly where and print some context
            for i, p in enumerate(doc.paragraphs):
                p_text = p.text.strip()
                if any(t.lower() in p_text.lower() for t in targets):
                    print(f"Para {i}: {p_text}")
                    # print next 20 paragraphs as well
                    for j in range(i+1, min(len(doc.paragraphs), i+40)):
                        print(f"  +{j-i}: {doc.paragraphs[j].text.strip()}")
                    break
    except Exception as e:
        print(f"Error reading {os.path.basename(f)}: {e}")
