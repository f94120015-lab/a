import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def create_comparison_report():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AMOK İNGİLİZCE ÖĞRENME UYGULAMASI\nİÇİNDEKİLER VE MÜFREDAT KARŞILAŞTIRMA RAPORU\n(Bölüm 1 - 18)")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 121) # Dark Blue

    doc.add_paragraph("\n")

    # Introduction
    p_intro = doc.add_paragraph()
    p_intro.add_run(
        "Bu rapor, AMOK İngilizce Öğrenme Uygulaması'nın veri tabanında (data.js) tanımlı ilk 18 ünitenin yapısı "
        "ile orijinal kitap içeriğindeki (amok düzenleme.docx) table of contents (İçindekiler) bölümünün "
        "birebir karşılaştırılması amacıyla hazırlanmıştır. Analiz sonucunda uygulamadaki eksik, fazla, yer değiştirmiş "
        "ve yeniden yapılandırılmış ders ve konular tespit edilmiştir."
    )
    p_intro.runs[0].font.name = 'Arial'
    p_intro.runs[0].font.size = Pt(11)

    # Summary Table Title
    h_summary = doc.add_paragraph()
    h_summary_run = h_summary.add_run("1. Genel Karşılaştırma ve Durum Tablosu")
    h_summary_run.font.name = 'Arial'
    h_summary_run.font.size = Pt(13)
    h_summary_run.font.bold = True
    h_summary_run.font.color.rgb = RGBColor(31, 78, 121)

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    
    headers = ["Kitap Bölümü (TOC)", "Uygulama Konusu (data.js)", "Durum", "Fark ve Açıklama"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        set_cell_background(cell, "1F4E79")
        run.font.color.rgb = RGBColor(255, 255, 255)

    comparison_data = [
        ("I. Yapılar (A-D) (Sf. 1-6)", "VI. Yapılar (A-D) (Sf. 1-6)", "Sıralama Kayması", "Kitapta en başta yer alırken, uygulamada 6. üniteye kaydırılmıştır."),
        ("II. A-F İsim ve Edat (Sf. 7-19)", "I. İsim ve Edat (Giriş + A-F) (Sf. 13-19)", "İçerik Değişikliği", "Uygulamada ek bir 'Giriş' dersi vardır. Kitaptaki II.A (Sf. 7) uygulamada E dersine (Sf. 13) taşınmıştır."),
        ("III. A-B Fiil ve Edat (Sf. 21-23)", "II. Fiil ve Edat (Sf. 21-23)", "Sıralama Kayması", "Yapılar'ın kaymasından ötürü Roma rakamı II. olarak adlandırılmıştır, içerik aynıdır."),
        ("IV. Özne + Fiil + Nesne (Sf. 32)", "VII. Özne + Fiil + Nesne (Sf. 32)", "Sıralama Kayması", "Roma rakamı VII. yapılmıştır, içerik birebir aynıdır."),
        ("V. There Yapısı (A-B) (Sf. 38-39)", "VIII. There Yapısı (A-B) (Sf. 38-39)", "Sıralama Kayması", "Roma rakamı VIII. yapılmıştır, içerik birebir aynıdır."),
        ("VI. Soru Strüktürleri (A-C) (Sf. 49-53)", "IX. Soru Yapıları (A-E) (Sf. 49-53)", "Sıralama & Harf Değişikliği", "Ders sıralaması tersine çevrilmiştir (Yardımcı fiil soruları başa alınmıştır). Alt harfler değişmiştir."),
        ("VII. Edilgen Strüktürü (Sf. 55)", "X. Edilgen Yapısı (A-D) (Sf. 55-62)", "Genişletilmiş İçerik", "Kitapta tek başlık olan edilgen yapısı, uygulamada 4 ayrı alt derse bölünmüştür."),
        ("VIII. Edilgen Mastarı (Sf. 63)", "XI. Edilgen Mastarı (A-F) (Sf. 63)", "Genişletilmiş İçerik", "Kitapta tek başlık olan edilgen mastarı, uygulamada 6 ayrı alt derse bölünmüştür."),
        ("IX. İsim Tamlaması (A-B) (Sf. 72-78)", "III. İsim Tamlaması (A-B) (Sf. 72-78)", "EKSİK İÇERİK", "Kitapta IX.B olan 'Sıfat + İsim (Sf. 78)' konusu uygulamada tamamen eksiktir."),
        ("X. Present Participle (A-B) (Sf. 81-84)", "IV. Present Participle (A-B) (Sf. 81-84)", "Sıralama Kayması", "Roma rakamı IV. yapılmıştır, içerik birebir aynıdır."),
        ("XI. Past Participle (A-B) (Sf. 85-88)", "V. Past Participle (A-B) (Sf. 85-88)", "Sıralama Kayması", "Roma rakamı V. yapılmıştır, içerik birebir aynıdır."),
        ("XII. Participle Takımları (A-B) (Sf. 88-90)", "XII. Participle Yapıları (A-C) (Sf. 88-90)", "Yeniden Yapılandırma", "Kitaptaki XII.A birleştirilmiş dersi uygulamada A ve B olarak ikiye bölünmüş, XII.B ise C yapılmıştır."),
        ("XIII. Mastar (A-B) (Sf. 95-98)", "XIII. Mastar (A-B) (Sf. 95-100)", "Eşleşiyor", "Roma rakamı aynı kalmıştır, sayfa numaralarında ufak farklar vardır."),
        ("XIV. Strüktürel It (A-B) (Sf. 98-103)", "XIV. Strüktürel It (A) (Sf. 103)", "EKSİK İÇERİK", "Kitaptaki XIV.B 'It + olmak + sıfat + mastar + nesnesi' uygulamada tamamen eksiktir."),
        ("XV. Maksat İçin Mastar (A-B) (Sf. 103-107)", "XV. Maksat ve Amaç (A-C) (Sf. 107-110)", "FAZLA İÇERİK", "Uygulamada akademik genişletilmiş cümleler içeren ekstra bir 'C. Maksat Yapıları' dersi eklenmiştir."),
        ("XVI. Fiil İsmi + Nesnesi (A-B) (Sf. 110-112)", "XVI. Fiil İsmi + Nesnesi (A) (Sf. 112)", "EKSİK İÇERİK", "Kitaptaki XVI.B 'Edattan sonra gelen fiil (+ nesnesi) (Sf. 112)' uygulamada tamamen eksiktir."),
        ("XVII. ...ing + Nesne (A-C) (Sf. 113-120)", "XVII. Edattan Sonra Fiil (A-C) (Sf. 113-120)", "Eşleşiyor", "Roma rakamı ve içerik eşleşiyor, başlık ismi kısaltılmıştır."),
        ("XVIII. Soru Kelimeli Mastar (Sf. 124)", "XVIII. Soru Kelimeli Mastar (A-C) (Sf. 124)", "Genişletilmiş İçerik", "Kitapta tek başlık olan konu uygulamada 3 ayrı alt ders halinde (120 soru) genişletilmiştir.")
    ]

    for item in comparison_data:
        row_cells = table.add_row().cells
        for col_idx, text in enumerate(item):
            row_cells[col_idx].text = text
            p = row_cells[col_idx].paragraphs[0]
            run = p.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)
            # Highlight missing content in red
            if "EKSİK" in text:
                run.font.bold = True
                run.font.color.rgb = RGBColor(192, 0, 0)
            elif "FAZLA" in text:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph("\n")

    # Section 2: Detailed Missing Content
    h_missing = doc.add_paragraph()
    h_missing_run = h_missing.add_run("2. Uygulamada Eksik Olan Bölüm ve Dersler")
    h_missing_run.font.name = 'Arial'
    h_missing_run.font.size = Pt(13)
    h_missing_run.font.bold = True
    h_missing_run.font.color.rgb = RGBColor(31, 78, 121)

    p_miss_desc = doc.add_paragraph(
        "Kitapta bulunmasına rağmen uygulamanın veri tabanında (data.js) yer almayan 3 ana konu/ders tespit edilmiştir:"
    )
    p_miss_desc.runs[0].font.name = 'Arial'
    p_miss_desc.runs[0].font.size = Pt(11)

    # Missing item 1
    m1 = doc.add_paragraph(style='List Bullet')
    m1_run = m1.add_run("Bölüm IX.B: Sıfat + İsim (Sayfa 78)\n")
    m1_run.font.bold = True
    m1_run.font.name = 'Arial'
    m1_run.font.size = Pt(11)
    m1_det = m1.add_run(
        "Açıklama: Kitaptaki 'İsim Tamlaması' ünitesinin ikinci dersidir. Uygulamada III. İsim Tamlaması ünitesi "
        "altında sadece 'İsim + isim' ve 'İsim + isim + isim' dersleri bulunmakta, 'Sıfat + isim' dersi bulunmamaktadır."
    )
    m1_det.font.name = 'Arial'
    m1_det.font.size = Pt(10.5)

    # Missing item 2
    m2 = doc.add_paragraph(style='List Bullet')
    m2_run = m2.add_run("Bölüm XIV.B: It + olmak + sıfat + mastar + nesnesi (Sayfa 103)\n")
    m2_run.font.bold = True
    m2_run.font.name = 'Arial'
    m2_run.font.size = Pt(11)
    m2_det = m2.add_run(
        "Açıklama: Kitapta 'Strüktürel It Cümlesinin Öznesi' ünitesinin ikinci alt başlığıdır. Uygulamada XIV. "
        "ünitede yalnızca tek bir ders (A. It + olmak + sıfat + mastar - Sf. 103) bulunmakta, nesneli olan B yapısı yer almamaktadır."
    )
    m2_det.font.name = 'Arial'
    m2_det.font.size = Pt(10.5)

    # Missing item 3
    m3 = doc.add_paragraph(style='List Bullet')
    m3_run = m3.add_run("Bölüm XVI.B: Edattan sonra gelen fiil (+ nesnesi) (Sayfa 112)\n")
    m3_run.font.bold = True
    m3_run.font.name = 'Arial'
    m3_run.font.size = Pt(11)
    m3_det = m3.add_run(
        "Açıklama: Kitaptaki 'Fiil ismi + nesnesi (özne olarak)' ünitesinin ikinci dersidir. "
        "Bu ders kapsamında yer alan 'This instrument is used for measuring...', 'There has been some discussion on controlling...' "
        "gibi edat sonrası gerund cümleleri uygulamaya aktarılmamıştır."
    )
    m3_det.font.name = 'Arial'
    m3_det.font.size = Pt(10.5)

    doc.add_paragraph("\n")

    # Section 3: Detailed Extra/Added Content
    h_extra = doc.add_paragraph()
    h_extra_run = h_extra.add_run("3. Uygulamada Fazla (Ekstra) Olan Bölüm ve Dersler")
    h_extra_run.font.name = 'Arial'
    h_extra_run.font.size = Pt(13)
    h_extra_run.font.bold = True
    h_extra_run.font.color.rgb = RGBColor(31, 78, 121)

    p_ext_desc = doc.add_paragraph(
        "Uygulamaya müfredat zenginleştirme amacıyla eklenen ve kitap içindekiler listesinde olmayan kısımlar:"
    )
    p_ext_desc.runs[0].font.name = 'Arial'
    p_ext_desc.runs[0].font.size = Pt(11)

    # Extra item 1
    e1 = doc.add_paragraph(style='List Bullet')
    e1_run = e1.add_run("Bölüm I (Giriş Dersi): İsim ve Edat Yapılarına Giriş (Sayfa 13)\n")
    e1_run.font.bold = True
    e1_run.font.name = 'Arial'
    e1_run.font.size = Pt(11)
    e1_det = e1.add_run(
        "Açıklama: Uygulamadaki I. İsim ve Edat Yapıları ünitesinde en başta yer alan 'Giriş' dersidir. "
        "Kitapta doğrudan A. İsim + of the + isim başlığı ile başlanır."
    )
    e1_det.font.name = 'Arial'
    e1_det.font.size = Pt(10.5)

    # Extra item 2
    e2 = doc.add_paragraph(style='List Bullet')
    e2_run = e2.add_run("Bölüm XV.C: Maksat Yapıları (Tam Genişletilmiş Akademik Örnekler)\n")
    e2_run.font.bold = True
    e2_run.font.name = 'Arial'
    e2_run.font.size = Pt(11)
    e2_det = e2.add_run(
        "Açıklama: Kitaptaki maksat yapılarının akademik genişletilmiş cümle versiyonlarını içeren "
        "ve uygulamaya özel olarak geliştirilip entegre edilen ders paketidir."
    )
    e2_det.font.name = 'Arial'
    e2_det.font.size = Pt(10.5)

    doc.add_paragraph("\n")

    # Section 4: Structure/Order differences
    h_structure = doc.add_paragraph()
    h_structure_run = h_structure.add_run("4. Sıralama ve Yapısal Farklılıklar")
    h_structure_run.font.name = 'Arial'
    h_structure_run.font.size = Pt(13)
    h_structure_run.font.bold = True
    h_structure_run.font.color.rgb = RGBColor(31, 78, 121)

    struct_items = [
        ("Ünite Roma Rakamlarındaki Kaymalar", "Kitapta en başta yer alan 'Yapılar' konusu, uygulamada 6. ünite yapıldığı için aradaki diğer tüm ünitelerin Roma rakamları ve indeksleri kaymıştır (örn. Kitap Bölüm IX -> Uygulama Bölüm III)."),
        ("Soru Yapıları Ders Sıralaması", "Kitapta soru kelimeli (Wh-) yapılar önce gelirken, uygulamada fiil başa getirilerek yapılan (Yes/No) soruları başa alınmıştır."),
        ("Derslerin Bölünerek Genişletilmesi", "Kitapta tek bir başlık altında bulunan 'Edilgen Strüktürü' (Bölüm VII) uygulamada 4 alt derse, 'Edilgen Mastarı' (Bölüm VIII) 6 alt derse ve son olarak 'Soru Kelimesinden Sonra Gelen Mastar' (Bölüm XVIII) 3 alt derse bölünerek alıştırma sayısı artırılmıştır."),
        ("Participle Takımlarının Yeniden Yapılandırılması", "Kitaptaki XII.A başlığı altındaki birleşik iki yapı, uygulamada A ve B olarak iki farklı ders şeklinde organize edilmiştir."),
        ("Sayfa Numaralarındaki Küçük Uyuşmazlıklar", "Bazı derslerin alt başlıklarında yer alan sayfa numaraları (örn. Giriş edat yapıları sf. 13 vs sf. 7) kitap ile uygulama arasında farklılıklar göstermektedir.")
    ]

    for s_title, s_desc in struct_items:
        sp = doc.add_paragraph(style='List Bullet')
        s_run = sp.add_run(s_title + ": ")
        s_run.font.bold = True
        s_run.font.name = 'Arial'
        s_run.font.size = Pt(11)
        
        s_desc_run = sp.add_run(s_desc)
        s_desc_run.font.name = 'Arial'
        s_desc_run.font.size = Pt(10.5)

    doc.add_paragraph("\n")

    # Conclusion
    h_concl = doc.add_paragraph()
    h_concl_run = h_concl.add_run("5. Sonuç ve Öneriler")
    h_concl_run.font.name = 'Arial'
    h_concl_run.font.size = Pt(13)
    h_concl_run.font.bold = True
    h_concl_run.font.color.rgb = RGBColor(31, 78, 121)

    p_concl = doc.add_paragraph(
        "Uygulamadaki ünitelerin sıralaması ve genişletilmiş yapısı pedagojik ve teknik açıdan oldukça zenginleştirilmiştir. "
        "Bununla birlikte, kitap ile tam uyumluluk sağlanması istenirse, eksik olan IX.B (Sıfat + İsim), XIV.B (Strüktürel It + Nesne) "
        "ve XVI.B (Edattan Sonra Fiil İsmi) derslerinin de yeni alıştırmalarla birlikte sisteme entegre edilmesi önerilmektedir."
    )
    p_concl.runs[0].font.name = 'Arial'
    p_concl.runs[0].font.size = Pt(11)

    # Save
    output_path = "/Users/faruknafizfazlioglu/Desktop/AMOK_Bolum_18_Karsilastirma_Raporu.docx"
    doc.save(output_path)
    print(f"Report saved successfully to: {output_path}")

if __name__ == "__main__":
    create_comparison_report()
