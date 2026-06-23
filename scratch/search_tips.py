import os
import glob
from docx import Document

desktop_dir = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_dir, "*.docx"))
docx_files += glob.glob(os.path.join(desktop_dir, "**/*.docx"), recursive=True)

target_phrases = [
    "emerging from the study",
    "Tercüme Kılavuzu",
    "Present Participle (V-ing) + Edat",
    "Tip: İsim + Present Participle",
    "Tip 1",
    "Tip 2",
    "Tip 3"
]

print(f"Searching in {len(docx_files)} docx files on Desktop...")

for fpath in docx_files:
    if "~$" in fpath:
        continue
    try:
        doc = Document(fpath)
        found = False
        for i, p in enumerate(doc.paragraphs):
            text = p.text
            for target in target_phrases:
                if target.lower() in text.lower():
                    print(f"\n--- Found '{target}' in {fpath} (Paragraph {i}) ---")
                    # Print 30 paragraphs around it to get full context
                    start = max(0, i - 5)
                    end = min(len(doc.paragraphs), i + 35)
                    for j in range(start, end):
                        print(f"P {j}: {doc.paragraphs[j].text.strip()}")
                    found = True
                    break
            if found:
                # Let's print some more if we want, or just continue to check other files
                found = False # Reset to find all occurrences
    except Exception as e:
        # print(f"Error reading {fpath}: {e}")
        pass
