from docx import Document
import os

files = [
    "/Users/faruknafizfazlioglu/Desktop/matching.docx",
    "/Users/faruknafizfazlioglu/Desktop/Amok kitabı-44444444.docx"
]

for f in files:
    if os.path.exists(f):
        doc = Document(f)
        print(f"{os.path.basename(f)}: paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}")
    else:
        print(f"Not found: {f}")
