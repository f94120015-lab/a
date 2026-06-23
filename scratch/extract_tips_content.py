import os
from docx import Document

fpath = "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx"
if os.path.exists(fpath):
    doc = Document(fpath)
    print(f"Total paragraphs in {fpath}: {len(doc.paragraphs)}")
    
    # We want to print from paragraph 5700 to 5800
    for j in range(5700, min(len(doc.paragraphs), 5900)):
        txt = doc.paragraphs[j].text.strip()
        if txt:
            print(f"P {j}: {txt}")
else:
    print("File not found")
