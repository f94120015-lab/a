import docx

doc = docx.Document("/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx")
print(f"Total paragraphs in amok düzenleme.docx: {len(doc.paragraphs)}")
for idx, p in enumerate(doc.paragraphs):
    if "after the ice melts" in p.text or "after the sun sets" in p.text:
        print(f"Match at {idx}: {p.text}")
        # Print next 120 paragraphs
        for i in range(idx, min(len(doc.paragraphs), idx + 150)):
            print(f"{i}: {doc.paragraphs[i].text}")
        break
