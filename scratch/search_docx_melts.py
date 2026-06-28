import docx
import os
import glob

files = glob.glob("/Users/faruknafizfazlioglu/Desktop/*.docx") + glob.glob("/Users/faruknafizfazlioglu/Desktop/**/*.docx")
for f in files:
    if os.path.basename(f).startswith("~$"):
        continue
    try:
        doc = docx.Document(f)
        print(f"File: {f}, Paras: {len(doc.paragraphs)}")
        for idx, p in enumerate(doc.paragraphs):
            if "after the ice melts" in p.text or "exposed to air" in p.text:
                print(f"MATCH at paragraph {idx} in {f}: {p.text}")
    except Exception as e:
        print(f"Error reading {f}: {e}")
