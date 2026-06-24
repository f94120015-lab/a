import os
import unicodedata
from docx import Document

filepath = "/Users/faruknafizfazlioglu/Desktop/soru örnekleri bölümler.docx"
filepath = unicodedata.normalize('NFD', filepath)

doc = Document(filepath)

# Let's print paragraphs 10 to 90 to see the beginning of Bölüm 7
for idx in range(10, 90):
    if idx < len(doc.paragraphs):
        print(f"P {idx}: {doc.paragraphs[idx].text.strip()}")
