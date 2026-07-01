import docx
import glob
import os

desktop_path = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_path, "*.docx"))

for filepath in docx_files:
    filename = os.path.basename(filepath)
    if "~$" in filename:
        continue
    try:
        doc = docx.Document(filepath)
        found = False
        lines = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                lines.append(text)
        
        for i, line in enumerate(lines):
            if "adverb + adjective" in line.lower() or "verb + adverb" in line.lower():
                print(f"\n========================================\nFOUND IN: {filename} (line {i})")
                # print 20 lines before and 40 lines after
                start = max(0, i - 10)
                end = min(len(lines), i + 40)
                for idx in range(start, end):
                    prefix = "--> " if idx == i else "    "
                    print(f"{prefix}{idx}: {lines[idx]}")
                found = True
        
        # Also check tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                if "adverb + adjective" in row_text.lower() or "verb + adverb" in row_text.lower():
                    print(f"\n========================================\nFOUND IN TABLE of {filename} (Table {t_idx}, Row {r_idx})")
                    print(row_text)
    except Exception as e:
        pass
