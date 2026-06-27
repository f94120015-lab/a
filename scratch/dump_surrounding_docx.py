import os
from docx import Document

def dump_surrounding():
    f = "/Users/faruknafizfazlioglu/Desktop/Amok Soru Yapılanlar.docx"
    if not os.path.exists(f):
        print(f"File not found: {f}")
        return
    try:
        doc = Document(f)
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        for idx in range(1170, min(1230, len(doc.paragraphs))):
            print(f"P {idx}: {doc.paragraphs[idx].text.strip()}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    dump_surrounding()
