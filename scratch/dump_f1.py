import os
from docx import Document
import unicodedata

f1 = "/Users/faruknafizfazlioglu/Desktop/Kübra Çalışma Dosyası. Diğerleri Sil - Kopya ilk hali.docx"
normalized = unicodedata.normalize('NFD', f1)
if not os.path.exists(normalized):
    normalized = unicodedata.normalize('NFC', f1)

try:
    doc = Document(normalized)
    with open("scratch/f1_paragraphs.txt", "w", encoding="utf-8") as out:
        for idx in range(1700, min(len(doc.paragraphs), 1950)):
            text = doc.paragraphs[idx].text.strip()
            out.write(f"P {idx}: {text}\n")
    print("Done writing to scratch/f1_paragraphs.txt")
except Exception as e:
    print(f"Error: {e}")
