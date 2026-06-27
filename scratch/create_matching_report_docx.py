import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # Page setup / margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("AMOK İNGİLİZCE ÖĞRENME UYGULAMASI\nKELİME EŞLEŞTİRME KÖK HALİ DÜZENLEME RAPORU")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(15)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 117, 89) # Pastel Dark Green

    doc.add_paragraph("\n")

    # Section 1: Introduction
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Kelime Eşleştirme Kök Hali Düzenlemesi Nedir?")
    h1_run.font.name = 'Arial'
    h1_run.font.size = Pt(13)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(46, 117, 89)

    p1 = doc.add_paragraph(
        "Kullanıcıların kelime öğrenimini kolaylaştırmak, kelimelerin cümle içindeki çekimli veya ekli halleri yerine "
        "sözlükteki yalın/kök halleriyle (master/infinitive) eşleştirilmesini sağlamak amacıyla uygulamadaki tüm "
        "ünitelerin eşleştirme alıştırmalarında genel bir düzenleme yapılmıştır. Yeni kurallar çerçevesinde:\n"
    )
    p1.runs[0].font.name = 'Arial'
    p1.runs[0].font.size = Pt(11)

    # Bullet points
    b1 = doc.add_paragraph(style='List Bullet')
    b1_run = b1.add_run("Eşleştirme kartlarında İngilizce kelimelerin daima yalın (V1) hali yer alır.")
    b1_run.font.name = 'Arial'
    b1_run.font.size = Pt(11)

    b2 = doc.add_paragraph(style='List Bullet')
    b2_run = b2.add_run("Türkçe karşılıklarda cümle içindeki çekimli ifadeler (örn. 'tanımlayabilir misiniz', 'tasarlamamın') yerine doğrudan sözlükteki kök/mastar halleri (örn. 'tanımlamak', 'tasarlamak') kullanılır.")
    b2_run.font.name = 'Arial'
    b2_run.font.size = Pt(11)

    b3 = doc.add_paragraph(style='List Bullet')
    b3_run = b3.add_run("Bu kural, bundan sonra eklenecek yeni ünitelerde de varsayılan kural olarak otomatik uygulanacaktır.")
    b3_run.font.name = 'Arial'
    b3_run.font.size = Pt(11)

    doc.add_paragraph("\n")

    # Section 2: Implementation Details
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Kuralın Teknik Uygulaması ve Mimari Değişiklikler")
    h2_run.font.name = 'Arial'
    h2_run.font.size = Pt(13)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(46, 117, 89)

    p2 = doc.add_paragraph(
        "Bu kuralı tüm ünitelerde tek seferde, dinamik ve geriye dönük olarak uygulayabilmek amacıyla "
        "kod altyapısında aşağıdaki mimari düzenlemeler gerçekleştirilmiştir:\n"
    )
    p2.runs[0].font.name = 'Arial'
    p2.runs[0].font.size = Pt(11)

    # Bullet list for implementation
    li1 = doc.add_paragraph(style='List Bullet')
    li1_run = li1.add_run("wordDictionary app.js dosyasından data.js dosyasının en üstüne taşınmıştır. Böylece veritabanı yüklenirken global olarak erişilebilir hale gelmiştir.")
    li1_run.font.name = 'Arial'
    li1_run.font.size = Pt(11)

    li2 = doc.add_paragraph(style='List Bullet')
    li2_run = li2.add_run("getUniqueMatchingPairs fonksiyonu, eşleştirme kelimelerini seçerken wordDictionary içindeki sözlük anlamlarını otomatik arar ve virgüle/taksim işaretine göre bölerek en temiz ilk sözlük karşılığını kartlara aktarır.")
    li2_run.font.name = 'Arial'
    li2_run.font.size = Pt(11)

    doc.add_paragraph("\n")

    # Section 3: Before vs. After Examples
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Yapılan Düzenlemelerden Örnekler (Eski vs. Yeni)")
    h3_run.font.name = 'Arial'
    h3_run.font.size = Pt(13)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(46, 117, 89)

    table = doc.add_table(rows=11, cols=4)
    table.style = 'Light Shading Accent 2'

    headers = ["İngilizce Kelime", "Eski Türkçe Kart Karşılığı", "Yeni Türkçe Kart Karşılığı", "Durum"]
    for i, head in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(head)
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(10)

    examples = [
        ("define", "tanımlayabilir misiniz", "tanımlamak", "Düzenlendi ✓"),
        ("present", "sunabilir misiniz", "sunmak", "Düzenlendi ✓"),
        ("restrict", "sınırlandıracak mısınız", "sınırlandırmak", "Düzenlendi ✓"),
        ("optimize", "optimize eder misiniz", "optimize etmek", "Düzenlendi ✓"),
        ("integrate", "entegre edebilir misiniz acaba", "entegre etmek", "Düzenlendi ✓"),
        ("shared", "paylaşmamın", "paylaşmak", "Düzenlendi ✓"),
        ("accessed", "erişmemin", "erişmek", "Düzenlendi ✓"),
        ("designed", "tasarlamamın", "tasarlamak", "Düzenlendi ✓"),
        ("conducted", "yürütmemin", "yürütmek", "Düzenlendi ✓"),
        ("inspected", "denetlememin", "denetlemek", "Düzenlendi ✓")
    ]

    for row_idx, data in enumerate(examples):
        row = table.rows[row_idx + 1].cells
        for col_idx, text in enumerate(data):
            row[col_idx].text = text

    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.5)

    doc.add_paragraph("\n")

    # Save document
    output_path = "/Users/faruknafizfazlioglu/Desktop/amok/Kelime_Eslestirme_Kok_Hal_Raporu.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == "__main__":
    create_report()
