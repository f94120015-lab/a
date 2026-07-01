import docx
import glob
import os

desktop_path = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_path, "*.docx"))

keywords = ["zarf", "sıfat", "adverb", "adjective", "noun", "isim"]

for filepath in docx_files:
    filename = os.path.basename(filepath)
    if "~$" in filename or "entegrasyon" in filename or "mufredat" in filename or "eksik" in filename:
        continue
    try:
        doc = docx.Document(filepath)
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if any(k in text.lower() for k in keywords):
                if len(text) > 10 and len(text) < 1000:
                    print(f"{filename} [P {i}]: {text[:150]}")
        
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                if any(k in row_text.lower() for k in keywords):
                    print(f"{filename} [Table {t_idx}, Row {r_idx}]: {row_text[:150]}")
    except Exception as e:
        pass
