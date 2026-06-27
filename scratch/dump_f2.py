import os
from docx import Document
import unicodedata

f2 = "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx"
normalized = unicodedata.normalize('NFD', f2)
if not os.path.exists(normalized):
    normalized = unicodedata.normalize('NFC', f2)

try:
    doc = Document(normalized)
    with open("scratch/f2_paragraphs.txt", "w", encoding="utf-8") as out:
        for idx in range(5080, min(len(doc.paragraphs), 5165)):
            text = doc.paragraphs[idx].text.strip()
            out.write(f"P {idx}: {text}\n")
    print("Done writing to scratch/f2_paragraphs.txt")
except Exception as e:
    print(f"Error: {e}")
