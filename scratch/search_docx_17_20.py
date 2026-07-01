import docx
import glob
import os

desktop_path = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_path, "*.docx"))

for filepath in docx_files:
    filename = os.path.basename(filepath)
    if "~$" in filename:
        continue
    try:
        doc = docx.Document(filepath)
        print(f"\n========================================\nFILE: {filename}")
        count = 0
        for p in doc.paragraphs:
            text = p.text
            if "17. Ders" in text or "20. Ders" in text or "17." in text or "20." in text or "Ders 17" in text or "Ders 20" in text:
                print(f"  P: {text}")
                count += 1
                if count > 20:
                    print("  ... (truncated)")
                    break
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                if "17" in row_text or "20" in row_text:
                    print(f"  TABLE ROW: {row_text}")
    except Exception as e:
        print(f"Error reading {filename}: {e}")
