import os
from docx import Document

fpath = "/Users/faruknafizfazlioglu/Desktop/dersler düzenleme.docx"
# Normalize unicode name
import unicodedata
f_norm = unicodedata.normalize('NFC', fpath)
if os.path.exists(f_norm):
    doc = Document(f_norm)
    print(f"Total paragraphs in {f_norm}: {len(doc.paragraphs)}")
    for j in range(95, min(len(doc.paragraphs), 160)):
        txt = doc.paragraphs[j].text.strip()
        if txt:
            print(f"P {j}: {txt}")
else:
    print("File not found")
