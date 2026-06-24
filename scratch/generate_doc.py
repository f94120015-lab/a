import os
import unicodedata
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load sentences directly from generate_js_arrays.py logic
# We can re-use the exact list we generated
components = [
    {
        "s1": "The data", "s2": "The newly collected empirical data", "v": "contradicts", "o1": "the theory.", "o2": "the long-standing theoretical model.",
        "s1_tr": "Veriler", "s2_tr": "Yeni toplanan deneysel veriler", "v_tr": "çelişir", "o1_tr": "teoriyle", "o2_tr": "uzun süredir var olan teorik modelle"
    },
    {
        "s1": "The context", "s2": "The broader socio-economic context", "v": "specifies", "o1": "the criteria.", "o2": "the strict qualitative selection criteria.",
        "s1_tr": "Bağlam", "s2_tr": "Daha geniş sosyo-ekonomik bağlam", "v_tr": "belirler", "o1_tr": "kriterleri", "o2_tr": "katı nitel seçim kriterlerini"
    },
    {
        "s1": "The sector", "s2": "The highly competitive dynamic sector", "v": "anticipates", "o1": "growth.", "o2": "significant annual financial growth.",
        "s1_tr": "Sektör", "s2_tr": "Son derece rekabetçi dinamik sektör", "v_tr": "öngörür", "o1_tr": "büyüme", "o2_tr": "yıllık önemli finansal büyüme"
    },
    {
        "s1": "Authorities", "s2": "Leading institutional authorities", "v": "advocate", "o1": "reform.", "o2": "comprehensive legislative tax reform.",
        "s1_tr": "Yetkililer", "s2_tr": "Önde gelen kurumsal yetkililer", "v_tr": "savunur", "o1_tr": "reformu", "o2_tr": "kapsamlı yasal vergi reformunu"
    },
    {
        "s1": "The dynamic", "s2": "This unpredictable economic dynamic", "v": "triggers", "o1": "reaction.", "o2": "a chain of negative physical reactions.",
        "s1_tr": "Dinamik", "s2_tr": "Bu öngörülemeyen ekonomik dinamik", "v_tr": "tetikler", "o1_tr": "tepkiyi", "o2_tr": "bir dizi olumsuz fiziksel tepkiyi"
    },
    {
        "s1": "Experts", "s2": "Independent technical experts", "v": "clarify", "o1": "the scope.", "o2": "the initial investigative project scope.",
        "s1_tr": "Uzmanlar", "s2_tr": "Bağımsız teknik uzmanlar", "v_tr": "açıklar", "o1_tr": "kapsamı", "o2_tr": "başlangıçtaki araştırma projesi kapsamını"
    },
    {
        "s1": "The process", "s2": "The continuous chemical process", "v": "induces", "o1": "stress.", "o2": "severe psychological and occupational stress.",
        "s1_tr": "Süreç", "s2_tr": "Sürekli kimyasal süreç", "v_tr": "yol açar", "o1_tr": "strese", "o2_tr": "ciddi psikolojik ve mesleki strese"
    },
    {
        "s1": "The anomaly", "s2": "The undetected structural anomaly", "v": "distorts", "o1": "results.", "o2": "the final statistical research results.",
        "s1_tr": "Anomali", "s2_tr": "Tespit edilemeyen yapısal anomali", "v_tr": "bozar", "o1_tr": "sonuçları", "o2_tr": "nihai istatistiksel araştırma sonuçlarını"
    },
    {
        "s1": "Media", "s2": "Mainstream digital media", "v": "manipulates", "o1": "perspective.", "o2": "public political and cultural perspective.",
        "s1_tr": "Medya", "s2_tr": "Ana akım dijital medya", "v_tr": "manipüle eder", "o1_tr": "bakış açısını", "o2_tr": "kamuoyunun siyasi ve kültürel bakış açısını"
    },
    {
        "s1": "The system", "s2": "The updated operational system", "v": "accommodates", "o1": "expansion.", "o2": "rapid regional infrastructure expansion.",
        "s1_tr": "Sistem", "s2_tr": "Güncellenmiş operasyonel sistem", "v_tr": "uyum sağlar", "o1_tr": "genişlemeye", "o2_tr": "hızlı bölgesel altyapı genişlemesine"
    },
    {
        "s1": "The protocol", "s2": "The revised security protocol", "v": "defines", "o1": "parameters.", "o2": "crucial technical system parameters.",
        "s1_tr": "Protokol", "s2_tr": "Gözden geçirilmiş güvenlik protokolü", "v_tr": "tanımlar", "o1_tr": "parametreleri", "o2_tr": "kritik teknik sistem parametrelerini"
    },
    {
        "s1": "The contract", "s2": "The legally binding contract", "v": "binds", "o1": "institutions.", "o2": "separate international research institutions.",
        "s1_tr": "Sözleşme", "s2_tr": "Yasal olarak bağlayıcı sözleşme", "v_tr": "bağlar", "o1_tr": "kurumları", "o2_tr": "ayrı uluslararası araştırma kurumlarını"
    },
    {
        "s1": "Analysts", "s2": "Senior financial analysts", "v": "inspect", "o1": "the framework.", "o2": "the entire underlying structural framework.",
        "s1_tr": "Analistler", "s2_tr": "Kıdemli finansal analistler", "v_tr": "inceler", "o1_tr": "çerçeveyi", "o2_tr": "tüm temel yapısal çerçeveyi"
    },
    {
        "s1": "The variable", "s2": "The primary independent variable", "v": "affects", "o1": "outcomes.", "o2": "excellent academic student outcomes.",
        "s1_tr": "Değişken", "s2_tr": "Birincil bağımsız değişken", "v_tr": "etkiler", "o1_tr": "sonuçları", "o2_tr": "mükemmel akademik öğrenci sonuçlarını"
    },
    {
        "s1": "The core", "s2": "The reinforced central core", "v": "stabilizes", "o1": "components.", "o2": "crucial internal device components.",
        "s1_tr": "Çekirdek", "s2_tr": "Güçlendirilmiş merkezi çekirdek", "v_tr": "stabilize eder", "o1_tr": "bileşenleri", "o2_tr": "kritik dahili cihaz bileşenlerini"
    },
    {
        "s1": "The graph", "s2": "The attached statistical graph", "v": "illustrates", "o1": "percentages.", "o2": "exact distribution and demographic percentages.",
        "s1_tr": "Grafik", "s2_tr": "Ekli istatistiksel grafik", "v_tr": "gösterir", "o1_tr": "yüzdeleri", "o2_tr": "kesin dağılım ve demografik yüzdeleri"
    },
    {
        "s1": "The policy", "s2": "The strict institutional policy", "v": "restricts", "o1": "access.", "o2": "unauthorized user network access.",
        "s1_tr": "Politika", "s2_tr": "Katı kurumsal politika", "v_tr": "kısıtlar", "o1_tr": "erişimi", "o2_tr": "yetkisiz kullanıcı ağ erişimini"
    },
    {
        "s1": "The finding", "s2": "The final scientific finding", "v": "validates", "o1": "hypotheses.", "o2": "alternative alternative scientific hypotheses.",
        "s1_tr": "Bulgu", "s2_tr": "Nihai bilimsel bulgu", "v_tr": "doğrular", "o1_tr": "hipotezleri", "o2_tr": "alternatif bilimsel hipotezleri"
    },
    {
        "s1": "The team", "s2": "The software development team", "v": "modifies", "o1": "modules.", "o2": "individual functional software modules.",
        "s1_tr": "Ekip", "s2_tr": "Yazılım geliştirme ekibi", "v_tr": "değiştirir", "o1_tr": "modülleri", "o2_tr": "bireysel fonksiyonel yazılım modüllerini"
    },
    {
        "s1": "The committee", "s2": "The ethics evaluation committee", "v": "evaluates", "o1": "feedback.", "o2": "detailed anonymous student feedback.",
        "s1_tr": "Komite", "s2_tr": "Etik değerlendirme komitesi", "v_tr": "değerlendirir", "o1_tr": "geri bildirimi", "o2_tr": "detaylı anonim öğrenci geri bildirimini"
    },
    {
        "s1": "The researcher", "s2": "The principal laboratory researcher", "v": "isolates", "o1": "variables.", "o2": "separate unstable chemical variables.",
        "s1_tr": "Araştırmacı", "s2_tr": "Baş laboratuvar araştırmacısı", "v_tr": "izole eder", "o1_tr": "değişkenleri", "o2_tr": "ayrı kararsız kimyasal değişkenleri"
    },
    {
        "s1": "Strategies", "s2": "Innovative corporate strategies", "v": "maximize", "o1": "efficiency.", "o2": "maximum annual manufacturing efficiency.",
        "s1_tr": "Stratejiler", "s2_tr": "Yenilikçi kurumsal stratejiler", "v_tr": "maksimize eder", "o1_tr": "verimliliği", "o2_tr": "maksimum yıllık üretim verimliliğini"
    },
    {
        "s1": "The script", "s2": "The automated background script", "v": "calculates", "o1": "ratios.", "o2": "complex mathematical data ratios.",
        "s1_tr": "Betik", "s2_tr": "Otomatik arka plan betiği", "v_tr": "hesaplar", "o1_tr": "oranları", "o2_tr": "karmaşık matematiksel veri oranlarını"
    },
    {
        "s1": "The audit", "s2": "The independent annual audit", "v": "exposes", "o1": "flaws.", "o2": "hidden organizational system flaws.",
        "s1_tr": "Denetim", "s2_tr": "Bağımsız yıllık denetim", "v_tr": "ortaya çıkarır", "o1_tr": "kusurları", "o2_tr": "gizli örgütsel sistem kusurlarını"
    },
    {
        "s1": "The shift", "s2": "The sudden paradigm shift", "v": "alters", "o1": "trends.", "o2": "global consumer behavior trends.",
        "s1_tr": "Değişim", "s2_tr": "Ani paradigma değişimi", "v_tr": "değiştirir", "o1_tr": "eğilimleri", "o2_tr": "küresel tüketici davranışı eğilimlerini"
    },
    {
        "s1": "The framework", "s2": "The advanced cryptographic framework", "v": "secures", "o1": "data.", "o2": "sensitive user information data.",
        "s1_tr": "Çerçeve", "s2_tr": "Gelişmiş kriptografik çerçeve", "v_tr": "güvenceye alır", "o1_tr": "verileri", "o2_tr": "hassas kullanıcı bilgileri verilerini"
    },
    {
        "s1": "The council", "s2": "The regional administrative council", "v": "suspended", "o1": "regulations.", "o2": "outdated environmental safety regulations.",
        "s1_tr": "Konsey", "s2_tr": "Bölgesel idari konsey", "v_tr": "askıya aldı", "o1_tr": "düzenlemeleri", "o2_tr": "güncelliğini yitirmiş çevresel güvenlik düzenlemelerini"
    },
    {
        "s1": "The board", "s2": "The executive internal board", "v": "terminated", "o1": "agreements.", "o2": "formal bilateral commercial agreements.",
        "s1_tr": "Yönetim kurulu", "s2_tr": "Yürütme iç kurulu", "v_tr": "feshetti", "o1_tr": "anlaşmaları", "o2_tr": "resmi ikili ticari anlaşmaları"
    },
    {
        "s1": "The ministry", "s2": "The national education ministry", "v": "conducted", "o1": "surveys.", "o2": "comprehensive regional educational surveys.",
        "s1_tr": "Bakanlık", "s2_tr": "Milli eğitim bakanlığı", "v_tr": "yürüttü", "o1_tr": "anketler", "o2_tr": "kapsamlı bölgesel eğitim anketleri"
    },
    {
        "s1": "The database", "s2": "The centralized cloud database", "v": "accumulates", "o1": "logs.", "o2": "detailed historical system logs.",
        "s1_tr": "Veritabanı", "s2_tr": "Merkezi bulut veritabanı", "v_tr": "biriktirir", "o1_tr": "günlükleri", "o2_tr": "detaylı geçmiş sistem günlüklerini"
    }
]

# Create word document
doc = Document()

# Set margins and page details
sections = doc.sections
for section in sections:
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

# Title
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("Bölüm 7: Özne - Geçişli Fiil + Nesne (SVO) Alıştırma Cümleleri")
title_run.font.size = Pt(18)
title_run.font.bold = True
title_run.font.name = 'Calibri'
title_run.font.color.rgb = RGBColor(31, 78, 121)

# Tercüme Kılavuzu
guide_p = doc.add_paragraph()
guide_p.paragraph_format.space_before = Pt(12)
guide_p.paragraph_format.space_after = Pt(12)
guide_p.paragraph_format.line_spacing = 1.15

guide_title_run = guide_p.add_run("Tercüme Kılavuzu:\n")
guide_title_run.font.bold = True
guide_title_run.font.size = Pt(12)
guide_title_run.font.name = 'Calibri'

guide_text_run = guide_p.add_run(
    "Akademik makalelerde özneler ve nesneler; sıfat tamlamaları ve sıfat cümleleri (relative clauses) gibi yan yapılarla çok fazla uzatılarak bazen 3-4 satırı bulabilir. "
    "Öğrenci eğer cümledeki \"Ana Özne ne?\", \"Geçişli Fiil ne?\" ve \"Bu fiilin etkilediği Nesne nerede?\" üçlüsünü (SVO) doğru tespit edebilirse, "
    "cümlenin etrafındaki tüm süsleri ve tümleçleri kolayca ayıklayarak ana fikri saniyeler içinde anlar."
)
guide_text_run.font.italic = True
guide_text_run.font.size = Pt(11)
guide_text_run.font.name = 'Calibri'

def add_set_header(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.name = 'Calibri'

def create_sentences_table(data_list):
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'İngilizce Cümle'
    hdr_cells[1].text = 'Türkçe Çeviri'
    hdr_cells[2].text = 'Hedef Kelime'
    hdr_cells[3].text = 'Boşluklu Hali'
    
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)

    for idx, item in enumerate(data_list):
        row_cells = table.add_row().cells
        row_cells[0].text = f"{idx+1}. {item['en']}"
        row_cells[1].text = item['tr']
        row_cells[2].text = f"{item['word']} ({item['trWord']})"
        row_cells[3].text = item['blank']
        
        # Apply font settings
        for cell in row_cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.05
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9.5)

# --- Set 1: Yalın SVO ---
add_set_header("1. Adım: Yalın (Kuyruksuz) Cümleler")
set1_data = []
for idx, comp in enumerate(components):
    set1_data.append({
        "en": f"{comp['s1']} {comp['v']} {comp['o1']}",
        "tr": f"{comp['s1_tr']} {comp['o1_tr']} {comp['v_tr']}.",
        "word": comp['v'],
        "trWord": comp['v_tr'],
        "blank": f"{comp['s1']} ___ {comp['o1']}"
    })
create_sentences_table(set1_data)

# --- Set 2: Genişletilmiş Özneli SVO ---
add_set_header("2. Adım: Genişletilmiş Özneli Cümleler")
set2_data = []
for idx, comp in enumerate(components):
    set2_data.append({
        "en": f"{comp['s2']} {comp['v']} {comp['o1']}",
        "tr": f"{comp['s2_tr']} {comp['o1_tr']} {comp['v_tr']}.",
        "word": comp['v'],
        "trWord": comp['v_tr'],
        "blank": f"{comp['s2']} ___ {comp['o1']}"
    })
create_sentences_table(set2_data)

# --- Set 3: Genişletilmiş Nesneli SVO ---
add_set_header("3. Adım: Genişletilmiş Nesneli Cümleler")
set3_data = []
for idx, comp in enumerate(components):
    o2_to_use = comp['o2']
    if comp['v'] == "validates":
        o2_to_use = "alternative alternative scientific hypotheses."
    if comp['v'] == "maximize":
        o2_to_use = "maximum annual manufacturing efficiency."
        
    set3_data.append({
        "en": f"{comp['s1']} {comp['v']} {o2_to_use}",
        "tr": f"{comp['s1_tr']} {comp['o2_tr']} {comp['v_tr']}.",
        "word": comp['v'],
        "trWord": comp['v_tr'],
        "blank": f"{comp['s1']} ___ {o2_to_use}"
    })
create_sentences_table(set3_data)

# --- Set 4: Tam Genişletilmiş SVO ---
add_set_header("4. Adım: Tam Genişletilmiş Akademik Cümleler (Full SVO)")
set4_data = []
for idx, comp in enumerate(components):
    o2_to_use = comp['o2']
    if comp['v'] == "validates":
        o2_to_use = "alternative alternative scientific hypotheses."
    if comp['v'] == "maximize":
        o2_to_use = "maximum annual manufacturing efficiency."
        
    set4_data.append({
        "en": f"{comp['s2']} {comp['v']} {o2_to_use}",
        "tr": f"{comp['s2_tr']} {comp['o2_tr']} {comp['v_tr']}.",
        "word": comp['v'],
        "trWord": comp['v_tr'],
        "blank": f"{comp['s2']} ___ {o2_to_use}"
    })
create_sentences_table(set4_data)

# Save document
dest_path = "/Users/faruknafizfazlioglu/Desktop/Bölüm 7 Alıştırmaları.docx"
doc.save(dest_path)
print(f"Document saved to {dest_path}")
