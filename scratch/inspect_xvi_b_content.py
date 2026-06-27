import os
from docx import Document
import unicodedata

def inspect_xvi_b():
    f = "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx"
    normalized = unicodedata.normalize('NFC', f)
    doc = Document(normalized)
    
    for idx in range(4800, 4861):
        if idx < len(doc.paragraphs):
            print(f"P {idx}: {doc.paragraphs[idx].text.strip()}")

if __name__ == "__main__":
    inspect_xvi_b()
