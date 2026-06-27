import os
import glob
from docx import Document

def search():
    paths = [
        "/Users/faruknafizfazlioglu/Desktop",
        "/Users/faruknafizfazlioglu/Downloads"
    ]
    for p in paths:
        if not os.path.exists(p):
            continue
        for root, dirs, files in os.walk(p):
            # Skip hidden dirs
            dirs[:] = [d for d in dirs if not d.startswith('.')]
            for f in files:
                if f.startswith('~$') or f.startswith('.'):
                    continue
                ext = os.path.splitext(f)[1].lower()
                full_path = os.path.join(root, f)
                try:
                    if ext in ['.txt', '.json', '.html', '.js', '.py', '.xml', '.csv']:
                        with open(full_path, "r", encoding="utf-8", errors="ignore") as file:
                            content = file.read()
                            if "would rather" in content.lower() and "bodrum" in content.lower():
                                print(f"FOUND IN TXT: {full_path}")
                    elif ext == '.docx':
                        doc = Document(full_path)
                        found = False
                        for p_idx, para in enumerate(doc.paragraphs):
                            if "would rather" in para.text.lower() and "bodrum" in para.text.lower():
                                found = True
                                break
                        if found:
                            print(f"FOUND IN DOCX: {full_path}")
                            # Let's print out the paragraph content around there
                            for i, para in enumerate(doc.paragraphs):
                                text = para.text.strip()
                                if "Yapı" in text or "would rather" in text.lower():
                                    print(f"  P {i}: {text}")
                except Exception as e:
                    pass

if __name__ == "__main__":
    search()
