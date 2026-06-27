import os
import glob
from docx import Document
import unicodedata

desktop_dir = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_dir, "*.docx"))

sentences = [
    "He is learning how to swim",
    "We need to decide what to say",
    "They don't know where to meet",
    "She asked the driver where to stop",
    "how to alter",
    "how to avoid",
    "mikrodan akademige",
    "bölüm 30"
]

for f in docx_files:
    if "~$" in f:
        continue
    f_norm = unicodedata.normalize('NFD', f)
    if not os.path.exists(f_norm):
        f_norm = unicodedata.normalize('NFC', f)
    if not os.path.exists(f_norm):
        continue
        
    try:
        doc = Document(f_norm)
        # Search paragraphs
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            text_lower = text.lower().replace(".", "").replace(",", "").strip()
            for s in sentences:
                s_clean = s.lower().replace(".", "").replace(",", "").strip()
                if s_clean in text_lower:
                    print(f"MATCH '{s}' in paragraph of {os.path.basename(f_norm)} at {i}: '{text}'")
                    
        # Search tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    text_lower = text.lower().replace(".", "").replace(",", "").strip()
                    for s in sentences:
                        s_clean = s.lower().replace(".", "").replace(",", "").strip()
                        if s_clean in text_lower:
                            print(f"MATCH '{s}' in table {t_idx} row {r_idx} cell {c_idx} of {os.path.basename(f_norm)}: '{text[:150]}'")
    except Exception as e:
        print(f"Error reading {os.path.basename(f_norm)}: {e}")
