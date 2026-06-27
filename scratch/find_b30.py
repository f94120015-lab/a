import os
import glob
from docx import Document
import unicodedata

desktop_dir = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_dir, "*.docx"))

targets = ["swim", "alter", "avoid", "bölüm 30", "wh- clause"]

print(f"Found {len(docx_files)} docx files.")

for f in docx_files:
    if "~$" in f:
        continue
    f_nfc = unicodedata.normalize('NFC', f)
    f_nfd = unicodedata.normalize('NFD', f)
    
    doc_path = None
    if os.path.exists(f):
        doc_path = f
    elif os.path.exists(f_nfc):
        doc_path = f_nfc
    elif os.path.exists(f_nfd):
        doc_path = f_nfd
        
    if not doc_path:
        print(f"Skipping (does not exist): {f}")
        continue
        
    print(f"Searching in: {os.path.basename(doc_path)}")
    try:
        doc = Document(doc_path)
        print(f"  Total paragraphs: {len(doc.paragraphs)}")
        found = False
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip().lower()
            for t in targets:
                if t in text:
                    print(f"  -> Match '{t}' at paragraph {i}: '{p.text[:60]}'...")
                    found = True
                    break
        if not found:
            print("  No matches.")
    except Exception as e:
        print(f"  Error reading {os.path.basename(doc_path)}: {e}")
