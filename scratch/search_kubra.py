import docx
import os

filepath = "/Users/faruknafizfazlioglu/Desktop/Kübra Çalışma Dosyası. Diğerleri Sil - Kopya ilk hali.docx"
if not os.path.exists(filepath):
    print("File not found")
else:
    doc = docx.Document(filepath)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    keywords = ["zarf", "sıfat", "fiil", "adverb", "adjective", "verb", "tamlama", "activity"]
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if any(k in text.lower() for k in keywords):
            print(f"P {i}: {text[:120]}")
