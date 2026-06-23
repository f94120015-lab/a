import os
import unicodedata
from docx import Document

desktop_dir = '/Users/faruknafizfazlioglu/Desktop'
target = "emerging from the study"

print("Starting scan...")
for root, dirs, files in os.walk(desktop_dir):
    # Skip common large directories to avoid getting stuck
    if any(p in root for p in ['node_modules', '.git', '.idea', '.vscode', 'Library', 'Cache', 'Local Settings']):
        continue
    for file in files:
        if file.endswith('.docx') and not file.startswith('~$'):
            fpath = os.path.join(root, file)
            f_norm = unicodedata.normalize('NFC', fpath)
            try:
                doc = Document(f_norm)
                for i, p in enumerate(doc.paragraphs):
                    if target.lower() in p.text.lower():
                        print(f"FOUND IN: {f_norm} at P {i}")
                        # Print surrounding text
                        start = max(0, i - 5)
                        end = min(len(doc.paragraphs), i + 100)
                        for j in range(start, end):
                            print(f"P {j}: {doc.paragraphs[j].text.strip()}")
                        break
            except Exception as e:
                # print(f"Error reading {f_norm}: {e}")
                pass
print("Scan completed.")
