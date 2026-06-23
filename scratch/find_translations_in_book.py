import os
from docx import Document
import unicodedata

fpath = "/Users/faruknafizfazlioglu/Desktop/amok düzenleme.docx"
f_norm = unicodedata.normalize('NFC', fpath)
if os.path.exists(f_norm):
    doc = Document(f_norm)
    print("Searching for translations in amok düzenleme.docx...")
    # Let's search for "the forests remaining in these areas"
    for i, p in enumerate(doc.paragraphs):
        if "the forests remaining in these areas" in p.text.lower():
            print(f"Found in paragraph {i}:")
            for j in range(max(0, i-2), min(len(doc.paragraphs), i+15)):
                print(f"  P {j}: {doc.paragraphs[j].text.strip()}")
            break
else:
    print("File not found")
