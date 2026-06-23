import os
import glob
from docx import Document

desktop_dir = "/Users/faruknafizfazlioglu/Desktop"
docx_files = glob.glob(os.path.join(desktop_dir, "*.docx"))

targets = [
    "forests remaining",
    "organisms living in the soil",
    "facts well-known",
    "compound composed of two elements",
    "exposed to the air"
]

for target in targets:
    print(f"\nSearching for '{target}'...")
    for fpath in docx_files:
        if "~$" in fpath:
            continue
        try:
            doc = Document(fpath)
            for i, p in enumerate(doc.paragraphs):
                if target.lower() in p.text.lower():
                    print(f"Found in {os.path.basename(fpath)} at P {i}:")
                    # print context
                    start = max(0, i - 2)
                    end = min(len(doc.paragraphs), i + 3)
                    for j in range(start, end):
                        print(f"  P {j}: {doc.paragraphs[j].text.strip()}")
        except Exception as e:
            print(f"Error reading {fpath}: {e}")
