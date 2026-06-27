import os
import glob
from docx import Document
import unicodedata

def search_toc():
    files = glob.glob("/Users/faruknafizfazlioglu/Desktop/*.docx")
    for f in files:
        if "~$" in f:
            continue
        try:
            doc = Document(f)
            found = False
            for i, p in enumerate(doc.paragraphs[:100]):
                text = p.text.strip()
                if "İÇİNDEKİLER" in text.upper():
                    print(f"FOUND IN: {os.path.basename(f)} at paragraph {i}: {text[:100]}")
                    found = True
                    break
            if not found:
                # also check tables
                for t_idx, table in enumerate(doc.tables[:5]):
                    for r_idx, row in enumerate(table.rows[:5]):
                        for c_idx, cell in enumerate(row.cells):
                            if "İÇİNDEKİLER" in cell.text.upper():
                                print(f"FOUND IN TABLE of {os.path.basename(f)}: table {t_idx}, row {r_idx}, cell {c_idx}")
                                found = True
                                break
                        if found:
                            break
                    if found:
                        break
        except Exception as e:
            print(f"Error reading {f}: {e}")

if __name__ == "__main__":
    search_toc()
